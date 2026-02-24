"""
LAA Uplift Narrative Generator
For Woodruff Billing Ltd - Internal Use Only
Converts solicitor PDF submissions into LAA-compliant narratives
"""

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk, font
import pdfplumber
import re
from datetime import datetime
import calendar
import os
from pathlib import Path
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

class PDFDataExtractor:
    """Extracts form data from solicitor's PDF submission"""
    
    def __init__(self):
        self.data = {
            'case_details': {},
            'panel_membership': [],
            'stage1_selections': [],
            'stage2_selections': [],
            'proposed_uplift': None
        }
    
    def extract_from_pdf(self, pdf_path):
        """Extract all relevant data from the PDF"""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                # Combine all pages into one text block
                full_text = ""
                for page in pdf.pages:
                    full_text += page.extract_text() + "\n"
                
                # Extract case details
                self._extract_case_details(full_text)
                
                # Extract panel membership
                self._extract_panel_membership(full_text)
                
                # Extract Stage 1 selections
                self._extract_stage1(full_text)
                
                # Extract Stage 2 selections
                self._extract_stage2(full_text)
                
                # Extract proposed uplift percentage
                self._extract_uplift_percentage(full_text)
                
                return self.data
                
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_case_details(self, text):
        """Extract fee earner, matter type, and case name"""
        # Look for Fee Earner
        fee_earner_match = re.search(r'Fee Earner:\s*(.+?)(?:\n|Matter Type:)', text)
        if fee_earner_match:
            self.data['case_details']['fee_earner'] = fee_earner_match.group(1).strip()
        
        # Look for Matter Type
        matter_type_match = re.search(r'Matter Type:\s*(.+?)(?:\n|Case)', text)
        if matter_type_match:
            self.data['case_details']['matter_type'] = matter_type_match.group(1).strip()
        
        # Look for Case/Matter Name
        case_name_match = re.search(r'Case / Matter Name:\s*(.+?)(?:\n|Panel)', text)
        if case_name_match:
            self.data['case_details']['case_name'] = case_name_match.group(1).strip()
    
    def _extract_panel_membership(self, text):
        """Extract panel membership information"""
        panel_section = re.search(r'Panel Membership\s*\n(.+?)(?:Stage 1:|$)', text, re.DOTALL)
        if panel_section:
            panel_text = panel_section.group(1)
            
            # Check for specific panels
            if 'Resolution Accredited Specialist Panel' in panel_text:
                self.data['panel_membership'].append('Resolution Accredited Specialist Panel')
            if 'Law Society Children Panel' in panel_text:
                self.data['panel_membership'].append('Law Society Children Panel')
            if 'Law Society Family Law Panel Advanced' in panel_text:
                self.data['panel_membership'].append('Law Society Family Law Panel Advanced')
    
    def _extract_stage1(self, text):
        """Extract Stage 1 threshold test selections and explanations"""
        stage1_section = re.search(r'Stage 1: Threshold Test Selections\s*\n(.+?)(?:Stage 2:|$)', text, re.DOTALL)
        if stage1_section:
            section_text = stage1_section.group(1)
            
            # Pattern to match items with explanations
            pattern = r'- (.+?)\s*\((.+?)\)\s*(?:Explanation:\s*(.+?)(?=\n-|\nStage|$))?'
            matches = re.findall(pattern, section_text, re.DOTALL)
            
            for match in matches:
                item = {
                    'selection': match[0].strip(),
                    'category': match[1].strip(),
                    'explanation': match[2].strip() if len(match) > 2 and match[2] else ""
                }
                # Clean up the explanation - remove extra whitespace
                if item['explanation']:
                    item['explanation'] = ' '.join(item['explanation'].split())
                self.data['stage1_selections'].append(item)
    
    def _extract_stage2(self, text):
        """Extract Stage 2 level of enhancement factors"""
        stage2_section = re.search(r'Stage 2: Level of Enhancement Factors\s*\n(.+?)(?:Solicitor\'s Proposed|$)', text, re.DOTALL)
        if stage2_section:
            section_text = stage2_section.group(1)
            
            # Pattern to match items with explanations
            pattern = r'- (.+?)\s*\((.+?)\)\s*(?:Explanation:\s*(.+?)(?=\n-|\nSolicitor|$))?'
            matches = re.findall(pattern, section_text, re.DOTALL)
            
            for match in matches:
                item = {
                    'selection': match[0].strip(),
                    'category': match[1].strip(),
                    'explanation': match[2].strip() if len(match) > 2 and match[2] else ""
                }
                # Clean up the explanation
                if item['explanation']:
                    item['explanation'] = ' '.join(item['explanation'].split())
                self.data['stage2_selections'].append(item)
    
    def _extract_uplift_percentage(self, text):
        """Extract the proposed uplift percentage"""
        uplift_match = re.search(r'Proposed Uplift Percentage:\s*(\d+)%?', text)
        if uplift_match:
            self.data['proposed_uplift'] = int(uplift_match.group(1))


class NarrativeGenerator:
    """Generates LAA-compliant narratives from extracted data"""
    
    def __init__(self):
        # LAA threshold mappings for proper referencing
        self.laa_references = {
            'exceptional_competence': {
                'section': 'CAG Section 12.8.1',
                'spec': 'Spec Para 6.13(a)',
                'description': 'exceptional competence, skill, or expertise'
            },
            'exceptional_speed': {
                'section': 'CAG Section 12.8.2',
                'spec': 'Spec Para 6.13(b)',
                'description': 'exceptional speed'
            },
            'exceptional_circumstances': {
                'section': 'CAG Section 12.8.3',
                'spec': 'Spec Para 6.13(c)',
                'description': 'exceptional circumstances, novelty, weight, or complexity'
            },
            'degree_responsibility': {
                'section': 'CAG Section 12.9(a)',
                'description': 'degree of responsibility accepted by the fee earner'
            },
            'care_speed_economy': {
                'section': 'CAG Section 12.9(b)',
                'description': 'care, speed, and economy'
            },
            'novelty_weight_complexity': {
                'section': 'CAG Section 12.9(c)',
                'description': 'novelty, weight, and complexity of the case'
            }
        }
    
    def generate_narrative(self, data):
        """Generate the complete narrative from extracted data"""
        narrative = []
        
        # Header
        narrative.append(("UPLIFT ENHANCEMENT JUSTIFICATION NARRATIVE", "title"))
        narrative.append(("", "normal"))
        
        # Case information
        narrative.append((f"Case: {data['case_details'].get('case_name', 'Not specified')}", "normal"))
        narrative.append((f"Matter Type: {data['case_details'].get('matter_type', 'Not specified')}", "normal"))
        narrative.append((f"Fee Earner: {data['case_details'].get('fee_earner', 'Not specified')}", "normal"))
        narrative.append(("", "normal"))
        
        # Introduction with uplift percentage
        uplift = data.get('proposed_uplift', 0)
        matter_type = data['case_details'].get('matter_type', 'referenced')
        case_name = data['case_details'].get('case_name', '[case name]')
        
        # Varied sentence structure for introduction
        intro_text = (f"An enhancement of {uplift}% is claimed on the {matter_type} work. "
                     f"This relates to the matter of {case_name}. "
                     f"The uplift is considered and justifiable. "
                     f"Exceptional factors detailed below meet the requirements in CPR 44.4(3). "
                     f"These factors also satisfy the Legal Aid Agency's Costs Assessment Guidance "
                     f"(Version 1a, 23 September 2024).")
        narrative.append((intro_text, "normal"))
        narrative.append(("", "normal"))
        
        # Panel Membership (if applicable)
        if data['panel_membership']:
            narrative.append(("PANEL MEMBERSHIP", "heading"))
            panel_text = "CAG Sections 12.20 to 12.23 provide for a minimum enhancement of 15%. "
            panel_text += "The fee earner holds membership of the following:"
            narrative.append((panel_text, "normal"))
            for panel in data['panel_membership']:
                narrative.append((f"• {panel}", "bullet"))
            narrative.append(("Work undertaken on this matter falls within the scope of this accreditation.", "normal"))
            narrative.append(("", "normal"))
        
        # Stage 1: Threshold Test
        if data['stage1_selections']:
            narrative.append(("STAGE 1: THRESHOLD TEST FOR ENHANCEMENT", "heading"))
            stage1_intro = "Work on this matter meets the threshold for enhancement. "
            stage1_intro += "Spec Para 6.13 and CAG Section 12.4 recognise the following exceptional factors."
            narrative.append((stage1_intro, "normal"))
            narrative.append(("", "normal"))
            
            # Group Stage 1 items by category
            categories = {}
            for item in data['stage1_selections']:
                category = self._get_category_type(item['category'])
                if category not in categories:
                    categories[category] = []
                categories[category].append(item)
            
            # Write each category
            for cat_type, items in categories.items():
                ref = self._get_reference_for_category(cat_type)
                if ref:
                    # Cleaner subheading without title case
                    desc = ref['description']
                    subheading = f"{ref['section']} and {ref['spec']} address {desc}:"
                    narrative.append((subheading, "subheading"))
                    for item in items:
                        narrative.append((f"• {item['selection']}", "bullet"))
                        if item['explanation'] and not self._is_generic_text(item['explanation']):
                            # Clean up explanation text
                            explanation = self._clean_explanation_text(item['explanation'])
                            if explanation:
                                narrative.append((f"  {explanation}", "explanation"))
                    narrative.append(("", "normal"))
        
        # Stage 2: Level of Enhancement
        if data['stage2_selections']:
            narrative.append(("STAGE 2: JUSTIFICATION FOR LEVEL OF ENHANCEMENT", "heading"))
            stage2_intro = f"A {uplift}% enhancement is justified under Spec Para 6.15. "
            stage2_intro += "CAG Sections 12.5 and 12.9 support this level based on the following factors."
            narrative.append((stage2_intro, "normal"))
            narrative.append(("", "normal"))
            
            # Group Stage 2 items by category
            categories = {}
            for item in data['stage2_selections']:
                category = self._get_stage2_category_type(item['category'])
                if category not in categories:
                    categories[category] = []
                categories[category].append(item)
            
            # Write each category
            for cat_type, items in categories.items():
                ref = self._get_stage2_reference(cat_type)
                if ref:
                    # Natural language subheading
                    desc = ref['description']
                    subheading = f"{ref['section']} considers {desc}:"
                    narrative.append((subheading, "subheading"))
                    for item in items:
                        narrative.append((f"• {item['selection']}", "bullet"))
                        if item['explanation'] and not self._is_generic_text(item['explanation']):
                            explanation = self._clean_explanation_text(item['explanation'])
                            if explanation:
                                narrative.append((f"  {explanation}", "explanation"))
                    narrative.append(("", "normal"))
        
        # Conclusion
        narrative.append(("CONCLUSION", "heading"))
        
        # More natural conclusion with varied sentence structure
        conclusion_parts = []
        conclusion_parts.append(f"The factors identified above rendered the work exceptionally demanding.")
        conclusion_parts.append(f"Both individually and together, these factors required exceptional skill.")
        conclusion_parts.append(f"The fee earner accepted responsibility beyond that normally expected at this level.")
        conclusion_parts.append(f"A {uplift}% enhancement is considered and justifiable under the LAA's assessment criteria.")
        conclusion_parts.append(f"Evidence supporting these assertions can be found within the case file.")
        conclusion_parts.append(f"Contemporaneous attendance notes provide additional verification.")
        
        conclusion_text = " ".join(conclusion_parts)
        narrative.append((conclusion_text, "normal"))
        
        return narrative
    
    def generate_word_document(self, data):
        """Generate a Word document with formatted narrative"""
        doc = Document()
        narrative = self.generate_narrative(data)
        
        for text, style in narrative:
            if style == "title":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(16)
                run.font.bold = True
            elif style == "heading":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(14)
                run.font.bold = True
                p.space_after = Pt(6)
            elif style == "subheading":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.italic = True
                p.space_after = Pt(3)
            elif style == "bullet":
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(text[2:])  # Remove bullet point as Word adds its own
            elif style == "explanation":
                p = doc.add_paragraph()
                p.left_indent = Pt(36)
                run = p.add_run(text.strip())
                run.font.italic = True
            else:  # normal
                if text:  # Only add paragraph if there's text
                    p = doc.add_paragraph(text)
        
        return doc
    
    def _is_generic_text(self, text):
        """Check if text is generic placeholder"""
        generic_phrases = [
            "LAA Enhancement Process Overview",
            "The LAA enhancement process involves two main stages",
            "Stage 1: A Threshold Test",
            "Stage 2: Determination of the Level of Enhancement",
            "This tool will guide you"
        ]
        return any(phrase in text for phrase in generic_phrases)
    
    def _clean_explanation_text(self, text):
        """Clean and format explanation text following style guide"""
        if not text or self._is_generic_text(text):
            return ""
        
        # Remove any semicolons, em dashes, en dashes
        text = text.replace(';', '.')
        text = text.replace('—', '.')
        text = text.replace('–', '.')
        
        # Remove overused AI vocabulary
        ai_words = {
            'robust': 'strong',
            'innovative': 'new',
            'delve into': 'examine',
            'delving': 'examining',
            'furthermore': '',
            'moreover': '',
            'therefore': '',
            'thus': '',
            'hence': '',
            'consequently': '',
            'in conclusion': '',
            'in summary': ''
        }
        
        text_lower = text.lower()
        for ai_word, replacement in ai_words.items():
            if ai_word in text_lower:
                # Case-insensitive replacement
                pattern = re.compile(re.escape(ai_word), re.IGNORECASE)
                text = pattern.sub(replacement, text)
        
        # Clean up double spaces and periods
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\.+', '.', text)
        text = text.strip()
        
        # Ensure it ends with a period
        if text and not text.endswith('.'):
            text += '.'
        
        return text
    
    def _format_date_uk(self, date_str):
        """Convert date to UK format (e.g., '11 December' or '30 December '25')"""
        # This would be used if dates are found in the text
        # For now, keeping dates as they appear in solicitor's input
        return date_str
    
    def _get_category_type(self, category_text):
        """Determine the category type from the text"""
        category_lower = category_text.lower()
        if 'competence' in category_lower or 'skill' in category_lower or 'expertise' in category_lower:
            return 'exceptional_competence'
        elif 'speed' in category_lower and 'exceptional' in category_lower:
            return 'exceptional_speed'
        elif 'circumstances' in category_lower or 'novelty' in category_lower or 'complexity' in category_lower:
            return 'exceptional_circumstances'
        return None
    
    def _get_stage2_category_type(self, category_text):
        """Determine Stage 2 category type"""
        category_lower = category_text.lower()
        if 'responsibility' in category_lower:
            return 'degree_responsibility'
        elif 'care' in category_lower or 'economy' in category_lower:
            return 'care_speed_economy'
        elif 'novelty' in category_lower or 'complexity' in category_lower:
            return 'novelty_weight_complexity'
        return None
    
    def _get_reference_for_category(self, cat_type):
        """Get LAA reference for a category"""
        return self.laa_references.get(cat_type)
    
    def _get_stage2_reference(self, cat_type):
        """Get Stage 2 LAA reference"""
        return self.laa_references.get(cat_type)


class UpliftNarrativeApp:
    """Main GUI Application"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("LAA Uplift Narrative Generator - Woodruff Billing Ltd")
        self.root.geometry("900x750")
        
        # Configure style
        self.root.configure(bg='#f0f0f0')
        
        # Create main frame
        main_frame = tk.Frame(root, bg='#f0f0f0', padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = tk.Label(main_frame, text="LAA Uplift Narrative Generator", 
                              font=('Arial', 16, 'bold'), bg='#f0f0f0')
        title_label.pack(pady=(0, 10))
        
        subtitle_label = tk.Label(main_frame, text="Internal Tool - Woodruff Billing Ltd", 
                                 font=('Arial', 10), bg='#f0f0f0', fg='#666')
        subtitle_label.pack(pady=(0, 20))
        
        # File selection frame
        file_frame = tk.Frame(main_frame, bg='#f0f0f0')
        file_frame.pack(fill=tk.X, pady=(0, 20))
        
        self.file_label = tk.Label(file_frame, text="No file selected", 
                                   bg='white', relief=tk.SUNKEN, padx=10, pady=5)
        self.file_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        browse_btn = tk.Button(file_frame, text="Browse PDF", 
                              command=self.browse_file, padx=20, pady=5)
        browse_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        # Process button
        self.process_btn = tk.Button(main_frame, text="Generate Narrative", 
                                     command=self.process_pdf, 
                                     font=('Arial', 12, 'bold'),
                                     bg='#4CAF50', fg='white', 
                                     padx=30, pady=10,
                                     state=tk.DISABLED,
                                     activeforeground='white',
                                     activebackground='#45a049')
        self.process_btn.pack(pady=(0, 20))
        
        # Output text area
        output_label = tk.Label(main_frame, text="Generated Narrative:", 
                               font=('Arial', 11, 'bold'), bg='#f0f0f0')
        output_label.pack(anchor=tk.W)
        
        # Text area with scrollbar
        text_frame = tk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 10))
        
        # Create custom font for the text widget
        text_font = font.Font(family='Calibri', size=11)
        self.output_text = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, font=text_font)
        self.output_text.pack(fill=tk.BOTH, expand=True)
        
        # Configure tags for formatting display
        self.output_text.tag_configure("title", font=('Calibri', 14, 'bold'), justify='center')
        self.output_text.tag_configure("heading", font=('Calibri', 12, 'bold'))
        self.output_text.tag_configure("subheading", font=('Calibri', 11, 'bold italic'))
        self.output_text.tag_configure("bullet", lmargin1=20, lmargin2=40)
        self.output_text.tag_configure("explanation", lmargin1=40, lmargin2=60, font=('Calibri', 10, 'italic'))
        
        # Buttons frame
        button_frame = tk.Frame(main_frame, bg='#f0f0f0')
        button_frame.pack(fill=tk.X)
        
        self.copy_btn = tk.Button(button_frame, text="Copy to Clipboard (for Word)", 
                                  command=self.copy_to_clipboard,
                                  bg='#1976D2', fg='#FFFFFF',
                                  padx=20, pady=8, state=tk.DISABLED,
                                  font=('Arial', 10, 'bold'),
                                  relief=tk.RAISED, bd=2,
                                  activeforeground='#FFFFFF',
                                  activebackground='#1565C0',
                                  disabledforeground='#CCCCCC')
        self.copy_btn.pack(side=tk.LEFT)
        
        self.save_btn = tk.Button(button_frame, text="Save as Text File", 
                                  command=self.save_narrative,
                                  bg='#757575', fg='#FFFFFF',
                                  padx=20, pady=8, state=tk.DISABLED,
                                  font=('Arial', 10),
                                  activeforeground='#FFFFFF',
                                  activebackground='#616161',
                                  disabledforeground='#CCCCCC')
        self.save_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        self.save_word_btn = tk.Button(button_frame, text="Save as Word Document", 
                                       command=self.save_as_word,
                                       bg='#0D47A1', fg='#FFFFFF',
                                       padx=20, pady=8, state=tk.DISABLED,
                                       font=('Arial', 10, 'bold'),
                                       activeforeground='#FFFFFF',
                                       activebackground='#0A3A7F',
                                       disabledforeground='#CCCCCC')
        self.save_word_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        clear_btn = tk.Button(button_frame, text="Clear", 
                             command=self.clear_all,
                             bg='#616161', fg='#FFFFFF',
                             padx=20, pady=8,
                             font=('Arial', 10),
                             activeforeground='#FFFFFF',
                             activebackground='#424242')
        clear_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        # Status bar
        self.status_label = tk.Label(main_frame, text="Ready", 
                                     bg='#ddd', relief=tk.SUNKEN, anchor=tk.W)
        self.status_label.pack(fill=tk.X, pady=(10, 0))
        
        # Store file path and data
        self.current_file = None
        self.current_narrative = None
        self.current_data = None
        self.formatted_narrative = None
    
    def browse_file(self):
        """Open file dialog to select PDF"""
        filename = filedialog.askopenfilename(
            title="Select Solicitor's PDF",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        
        if filename:
            self.current_file = filename
            self.file_label.config(text=os.path.basename(filename))
            self.process_btn.config(state=tk.NORMAL)
            self.status_label.config(text=f"File loaded: {os.path.basename(filename)}")
    
    def process_pdf(self):
        """Process the selected PDF and generate narrative"""
        if not self.current_file:
            messagebox.showerror("Error", "Please select a PDF file first")
            return
        
        try:
            # Update status
            self.status_label.config(text="Processing PDF...")
            self.root.update()
            
            # Extract data from PDF
            extractor = PDFDataExtractor()
            self.current_data = extractor.extract_from_pdf(self.current_file)
            
            # Generate narrative
            generator = NarrativeGenerator()
            self.formatted_narrative = generator.generate_narrative(self.current_data)
            
            # Clear and display narrative with formatting
            self.output_text.delete(1.0, tk.END)
            
            # Display with formatting tags
            for text, style in self.formatted_narrative:
                if text:  # Only add if there's text
                    self.output_text.insert(tk.END, text + "\n", style)
            
            # Generate plain text version for saving
            plain_lines = []
            for text, style in self.formatted_narrative:
                if style == "title":
                    plain_lines.append(text)
                    plain_lines.append("=" * 50)
                elif style == "heading":
                    plain_lines.append("")
                    plain_lines.append(text)
                    plain_lines.append("-" * 30)
                elif style == "subheading":
                    plain_lines.append(text)
                else:
                    plain_lines.append(text)
            
            self.current_narrative = "\n".join(plain_lines)
            
            # Enable buttons
            self.copy_btn.config(state=tk.NORMAL)
            self.save_btn.config(state=tk.NORMAL)
            self.save_word_btn.config(state=tk.NORMAL)
            
            # Update status
            self.status_label.config(text="Narrative generated successfully")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process PDF:\n{str(e)}")
            self.status_label.config(text="Error processing file")
    
    def copy_to_clipboard(self):
        """Copy formatted narrative to clipboard for Word"""
        if not self.current_data:
            messagebox.showerror("Error", "No narrative to copy")
            return
        
        try:
            # Create a temporary Word document
            generator = NarrativeGenerator()
            doc = generator.generate_word_document(self.current_data)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                doc.save(tmp.name)
                temp_path = tmp.name
            
            # Copy the content using Word automation if available
            # For now, we'll copy the formatted text in a way that preserves structure
            formatted_text = ""
            for text, style in self.formatted_narrative:
                if style == "title":
                    formatted_text += f"**{text}**\n\n"
                elif style == "heading":
                    formatted_text += f"\n**{text}**\n"
                elif style == "subheading":
                    formatted_text += f"*{text}*\n"
                elif style == "bullet":
                    formatted_text += f"{text}\n"
                elif style == "explanation":
                    formatted_text += f"    {text}\n"
                elif text:  # normal text
                    formatted_text += f"{text}\n"
            
            # Copy to clipboard
            self.root.clipboard_clear()
            self.root.clipboard_append(formatted_text)
            
            # Clean up temp file
            os.unlink(temp_path)
            
            messagebox.showinfo("Success", "Narrative copied to clipboard!\n\nPaste into Word for formatted text.")
            self.status_label.config(text="Copied to clipboard")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to copy to clipboard:\n{str(e)}")
    
    def save_narrative(self):
        """Save the generated narrative to a text file"""
        if not self.current_narrative:
            messagebox.showerror("Error", "No narrative to save")
            return
        
        # Generate default filename with UK date format
        now = datetime.now()
        date_str = now.strftime("%d_%B_%Y")
        default_name = f"uplift_narrative_{date_str}.txt"
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            initialfile=default_name
        )
        
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(self.current_narrative)
                
                messagebox.showinfo("Success", f"Narrative saved to:\n{filename}")
                self.status_label.config(text=f"Saved: {os.path.basename(filename)}")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file:\n{str(e)}")
    
    def save_as_word(self):
        """Save the narrative as a Word document"""
        if not self.current_data:
            messagebox.showerror("Error", "No narrative to save")
            return
        
        # Generate default filename with UK date format
        now = datetime.now()
        date_str = now.strftime("%d_%B_%Y")
        default_name = f"uplift_narrative_{date_str}.docx"
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=default_name
        )
        
        if filename:
            try:
                generator = NarrativeGenerator()
                doc = generator.generate_word_document(self.current_data)
                doc.save(filename)
                
                messagebox.showinfo("Success", f"Word document saved to:\n{filename}")
                self.status_label.config(text=f"Saved: {os.path.basename(filename)}")
                
                # Ask if user wants to open the document
                if messagebox.askyesno("Open Document", "Would you like to open the Word document now?"):
                    os.startfile(filename)
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save Word document:\n{str(e)}")
    
    def clear_all(self):
        """Clear all fields and reset"""
        self.current_file = None
        self.current_narrative = None
        self.current_data = None
        self.formatted_narrative = None
        self.file_label.config(text="No file selected")
        self.output_text.delete(1.0, tk.END)
        self.process_btn.config(state=tk.DISABLED)
        self.copy_btn.config(state=tk.DISABLED)
        self.save_btn.config(state=tk.DISABLED)
        self.save_word_btn.config(state=tk.DISABLED)
        self.status_label.config(text="Ready")


def main():
    """Main entry point"""
    root = tk.Tk()
    app = UpliftNarrativeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()