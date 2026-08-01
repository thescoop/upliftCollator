"""Verify the Markdown → Word conversion of the polished narrative.

The .docx is the file that actually goes to the LAA, while ``citation-check.txt``
certifies the Markdown. The weight of this suite is therefore on the one
question that gap creates: **can the Word document ever carry fewer authorities
than the document the check signed off?**

The fixture is the genuine output of a full pipeline run over the synthetic
``sample.pdf`` — real model prose with its real quirks (``*   `` bullets, bold
run-in leads, an italicised case name), not Markdown invented to suit the
parser. Nothing here touches the network, Word, or a real case.
"""

import sys
import unittest
from pathlib import Path
from unittest import mock

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import checks  # noqa: E402
import docx_writer  # noqa: E402
from docx_writer import Span  # noqa: E402

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "sample_polished.md"
POLISHED = FIXTURE.read_text(encoding="utf-8")


class TestBlockParsing(unittest.TestCase):
    def test_heading_levels(self):
        blocks = docx_writer.parse_blocks("## Introduction\n\n### Sub")
        self.assertEqual([(b.kind, b.text, b.level) for b in blocks],
                         [("heading", "Introduction", 2), ("heading", "Sub", 3)])

    def test_the_models_actual_bullet_form(self):
        """The model emits ``*   `` — asterisk, three spaces — not ``- ``."""
        blocks = docx_writer.parse_blocks("*   **Lead:** body text")
        self.assertEqual(blocks[0].kind, "bullet")
        self.assertEqual(blocks[0].text, "**Lead:** body text")

    def test_soft_wrapped_prose_becomes_one_paragraph(self):
        """Markdown soft wraps must not become hard line breaks in Word."""
        blocks = docx_writer.parse_blocks("One line\ncontinues here.\n\nSecond.")
        self.assertEqual([b.text for b in blocks],
                         ["One line continues here.", "Second."])

    def test_italic_line_is_not_mistaken_for_a_bullet(self):
        """``*Re B-S ...*`` opens with an asterisk but is emphasis, not a list."""
        blocks = docx_writer.parse_blocks("*Re B-S (Children) [2013] EWCA Civ 1146*")
        self.assertEqual(blocks[0].kind, "para")

    def test_numbered_and_quoted_blocks(self):
        blocks = docx_writer.parse_blocks("1. First\n\n> quoted line")
        self.assertEqual([b.kind for b in blocks], ["numbered", "quote"])

    def test_horizontal_rule_dropped_carries_no_text(self):
        blocks = docx_writer.parse_blocks("Above\n\n---\n\nBelow")
        self.assertEqual([b.text for b in blocks], ["Above", "Below"])


class TestInlineParsing(unittest.TestCase):
    def test_bold_italic_and_plain(self):
        spans = docx_writer.parse_inline("plain **bold** and *italic* end")
        self.assertEqual(
            [(s.text, s.bold, s.italic) for s in spans],
            [("plain ", False, False), ("bold", True, False), (" and ", False, False),
             ("italic", False, True), (" end", False, False)],
        )

    def test_bold_wins_over_italic_on_the_longer_marker(self):
        spans = docx_writer.parse_inline("**Lead:** rest")
        self.assertEqual((spans[0].text, spans[0].bold, spans[0].italic),
                         ("Lead:", True, False))

    def test_triple_marker_is_bold_and_italic(self):
        spans = docx_writer.parse_inline("***both***")
        self.assertTrue(spans[0].bold and spans[0].italic)

    def test_unmatched_marker_survives_as_literal_text(self):
        """Losing a stray asterisk is harmless; losing the words is not."""
        spans = docx_writer.parse_inline("2 * 3 = 6 and a lone ** here")
        self.assertEqual("".join(s.text for s in spans), "2 * 3 = 6 and a lone ** here")

    def test_intra_word_asterisks_are_not_emphasis(self):
        spans = docx_writer.parse_inline("cost*rate*hours")
        self.assertEqual("".join(s.text for s in spans), "cost*rate*hours")
        self.assertFalse(any(s.italic for s in spans))


class TestHeadingDropRule(unittest.TestCase):
    """Only citation-free scaffolding may be dropped. This is the guard that
    stops the hand-formatting mistake this module exists to prevent."""

    def test_pure_scaffolding_is_droppable(self):
        self.assertTrue(docx_writer._heading_is_droppable("Introduction"))
        self.assertTrue(docx_writer._heading_is_droppable("Conclusion"))

    def test_a_heading_carrying_a_citation_is_never_droppable(self):
        heading = ("Threshold Test (Qualifying for Enhancement - "
                   "Spec Para 6.13 / CAG Section 12.4)")
        self.assertFalse(docx_writer._heading_is_droppable(heading))

    def test_scaffolding_name_plus_citation_is_still_kept(self):
        """The name is not the test — the citation is."""
        self.assertFalse(
            docx_writer._heading_is_droppable("Introduction (CPR 44.4(3))")
        )

    def test_an_unrecognised_heading_is_kept(self):
        self.assertFalse(docx_writer._heading_is_droppable("Panel Membership"))


class TestRenderedDocument(unittest.TestCase):
    def setUp(self):
        self.doc, self.result = docx_writer.build_document(POLISHED)
        self.paras = self.doc.paragraphs

    def test_opens_with_the_bill_narrative_section_heading(self):
        self.assertEqual(self.paras[0].text, docx_writer.SECTION_HEADING)
        self.assertTrue(self.paras[0].runs[0].bold)

    def test_criteria_are_bulleted(self):
        """Five ticked criteria in the fixture, five bullets in the document."""
        bullets = [p for p in self.paras if p.style.name == "List Bullet"]
        self.assertEqual(len(bullets), 5)

    def test_each_bullet_opens_with_a_bold_run_in_lead(self):
        for para in (p for p in self.paras if p.style.name == "List Bullet"):
            self.assertTrue(para.runs[0].bold, f"not bold: {para.text[:40]!r}")
            self.assertFalse(para.runs[-1].bold,
                             "the whole bullet is bold; the evidence must be plain")

    def test_headings_are_demoted_to_bold_not_deleted(self):
        self.assertIn(
            "Threshold Test (Qualifying for Enhancement - Spec Para 6.13 / "
            "CAG Section 12.4)",
            self.result.demoted_headings,
        )
        rendered = [p.text for p in self.paras]
        self.assertTrue(any("Threshold Test" in t for t in rendered))

    def test_only_the_structural_headings_were_dropped(self):
        self.assertEqual(sorted(self.result.dropped_headings),
                         ["Conclusion", "Introduction"])

    def test_the_case_name_stays_italic(self):
        """Italicising a case name is a legal drafting convention, not noise."""
        italic = [r.text for p in self.paras for r in p.runs if r.italic]
        self.assertTrue(any("Re B-S" in t for t in italic), italic)

    def test_no_markdown_markers_survive_into_the_document(self):
        text = docx_writer.document_text(self.doc)
        self.assertNotIn("**", text)
        self.assertNotIn("## ", text)

    def test_every_prose_paragraph_survives(self):
        """Nothing but a dropped structural heading may go missing."""
        text = docx_writer.document_text(self.doc)
        for block in docx_writer.parse_blocks(POLISHED):
            if block.kind == "heading" and docx_writer._heading_is_droppable(block.text):
                continue
            plain = "".join(s.text for s in docx_writer.parse_inline(block.text))
            self.assertIn(plain, text, f"lost: {plain[:60]!r}")

    def test_typography_matches_the_bill_narrative(self):
        normal = self.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Arial Nova Cond Light")
        self.assertEqual(normal.font.size, docx_writer.BODY_SIZE)
        section = self.doc.sections[0]
        self.assertEqual(section.page_width, docx_writer.PAGE_WIDTH)
        self.assertEqual(section.left_margin, docx_writer.MARGIN)


class TestCitationFidelity(unittest.TestCase):
    """The reason this module exists."""

    def test_the_docx_carries_exactly_the_markdown_citations(self):
        doc, _ = docx_writer.build_document(POLISHED)
        self.assertEqual(
            sorted(checks._normalise(c) for c in checks.extract_citations(POLISHED)),
            sorted(checks._normalise(c)
                   for c in checks.extract_citations(docx_writer.document_text(doc))),
        )

    def test_the_fixture_really_does_hide_a_citation_in_a_heading(self):
        """Guards the guard: if the model stopped putting citations in headings
        this suite would pass while testing nothing."""
        headings = [b.text for b in docx_writer.parse_blocks(POLISHED)
                    if b.kind == "heading"]
        self.assertTrue(
            any(checks.extract_citations(h) for h in headings),
            "no heading carries a citation — the drop rule is untested",
        )

    def test_a_citation_split_across_a_bold_boundary_is_still_found(self):
        markdown = "**Care, speed, and economy (CAG** 12.9(b)): body text."
        doc, _ = docx_writer.build_document(markdown)
        self.assertIn("CAG 12.9(b)",
                      checks.extract_citations(docx_writer.document_text(doc)))

    def test_write_refuses_to_save_a_document_that_lost_a_citation(self):
        real = docx_writer.build_document

        def lossy(markdown):
            return real("Body text with no authority in it at all.")

        with mock.patch.object(docx_writer, "build_document", lossy):
            with self.assertRaises(docx_writer.NarrativeConversionError) as ctx:
                docx_writer.write_docx(POLISHED, Path("/dev/null"))
        self.assertIn("Missing from the .docx", str(ctx.exception))

    def test_write_refuses_when_a_citation_appears_only_in_the_docx(self):
        real = docx_writer.build_document

        def inventive(markdown):
            return real(markdown + "\n\nAlso CAG Section 99.9 applies.")

        with mock.patch.object(docx_writer, "build_document", inventive):
            with self.assertRaises(docx_writer.NarrativeConversionError) as ctx:
                docx_writer.write_docx(POLISHED, Path("/dev/null"))
        self.assertIn("only in the .docx", str(ctx.exception))

    def test_repeated_citations_are_counted_not_deduplicated(self):
        markdown = "First CPR 44.4(3) here.\n\nSecond CPR 44.4(3) there."
        doc, _ = docx_writer.build_document(markdown)
        found = checks.extract_citations(docx_writer.document_text(doc))
        self.assertEqual(len([c for c in found if "44.4" in c]), 2)


class TestWriteToDisk(unittest.TestCase):
    def test_round_trip_through_a_real_file(self):
        import tempfile

        from docx import Document as OpenDocument

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "narrative-polished.docx"
            outcome = docx_writer.write_docx(POLISHED, path)
            self.assertTrue(path.is_file())
            self.assertEqual(outcome.path, path)
            reopened = OpenDocument(str(path))
            self.assertEqual(reopened.paragraphs[0].text,
                             docx_writer.SECTION_HEADING)
            self.assertEqual(
                sorted(checks._normalise(c) for c in outcome.citations),
                sorted(checks._normalise(c)
                       for c in checks.extract_citations(
                           docx_writer.document_text(reopened))),
            )


if __name__ == "__main__":
    unittest.main()
