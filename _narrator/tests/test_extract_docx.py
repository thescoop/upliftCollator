"""The .docx extraction contract, end to end and under attack.

The fixtures are built by ``build_docx_fixture.js`` with the real browser
generator and labels read live from ``content-data.js`` — so these tests hold
the actual shipped round trip, not a Python imitation of it. Regenerate with::

    node _narrator/tests/build_docx_fixture.js

The adversarial cases matter most. ``nasty.docx`` carries current shaded-heading
text, a bare copy of the Stage 1 composite heading's first line, coded item rows,
detail rows and a fake uplift inside one explanation paragraph. Every line must
stay inert because the extractor treats that whole paragraph as opaque.
"""

from __future__ import annotations

import json
import shutil
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document

import extract
import extract_docx
from templates import load_content_data

FIXTURES = Path(__file__).resolve().parent / "fixtures"
SAMPLE = FIXTURES / "sample.docx"
DEEMED = FIXTURES / "deemed.docx"
NASTY = FIXTURES / "nasty.docx"


def _current_label(key: str) -> str:
    for block in load_content_data()["question_blocks"]:
        for chk in block.get("checkboxes", []):
            if chk["key"] == key:
                return chk["label"]
    raise KeyError(key)


def _current_checkbox(key: str) -> tuple[dict, dict]:
    for block in load_content_data()["question_blocks"]:
        for checkbox in block.get("checkboxes", []):
            if checkbox["key"] == key:
                return block, checkbox
    raise KeyError(key)


def _coded_row(key: str) -> str:
    _block, checkbox = _current_checkbox(key)
    return f"{checkbox['code']}\t{checkbox['label']}"


def _rewrite_paragraph(doc_path: Path, old_text: str, new_text: str) -> Path:
    """Copy a fixture with one paragraph's text replaced — a document someone
    edited in Word, which is the only way a Collator docx changes at all."""
    tmp = Path(tempfile.mkdtemp()) / doc_path.name
    shutil.copy(doc_path, tmp)
    doc = Document(str(tmp))
    for para in doc.paragraphs:
        if para.text == old_text:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text
            break
    else:
        raise AssertionError(f"paragraph not found: {old_text!r}")
    doc.save(str(tmp))
    return tmp


class TestSampleRoundTrip(unittest.TestCase):
    """The ordinary route: everything written is everything read."""

    @classmethod
    def setUpClass(cls):
        cls.data = extract.extract_formdata(SAMPLE)

    def test_case_details_come_back_whole(self):
        self.assertEqual(self.data["caseDetails"], {
            "feeEarnerName": "Jane Doe",
            "matterType": "Public Law Children",
            "caseMatterName": "Re X (Local Authority care proceedings)",
            "courtLevel": "County Court",
        })

    def test_all_sections_resolve_to_their_keys(self):
        self.assertEqual(sorted(self.data["panelMembership"]),
                         ["panel_membership_children", "panel_membership_resolution"])
        self.assertEqual(sorted(self.data["stage1"]),
                         ["s1_circ_legal_issues", "s1_cse_detailed_knowledge",
                          "s1_cse_marshalling_evidence"])
        self.assertEqual(sorted(self.data["stage2"]),
                         ["s2_care_vulnerable_client", "s2_resp_no_counsel_drafting"])

    def test_labels_are_the_current_wording(self):
        """A v1.13 document carries v1.13 labels — never legacy aliases."""
        for section in ("stage1", "stage2"):
            for key, entry in self.data[section].items():
                with self.subTest(key=key):
                    self.assertEqual(entry["label"], _current_label(key))

    def test_codes_are_reconstructed_from_content_data(self):
        for section in ("stage1", "stage2"):
            for key, entry in self.data[section].items():
                with self.subTest(key=key):
                    _block, checkbox = _current_checkbox(key)
                    self.assertEqual(entry["code"], checkbox["code"])

    def test_explanations_round_trip_exactly(self):
        exp = self.data["stage2"]["s2_care_vulnerable_client"]["explanation"]
        self.assertTrue(exp.startswith("Client has diagnosed PTSD"))
        self.assertTrue(exp.endswith("intermediary to the FHDRA."))

    def test_scalars(self):
        self.assertEqual(self.data["finalUpliftPercent"], "75")
        self.assertTrue(self.data["evidenceOnFileConfirmed"])
        self.assertFalse(self.data["thresholdDeemed"])
        self.assertNotIn("unrecognised", self.data)

    def test_category_titles_are_the_block_titles(self):
        self.assertEqual(
            self.data["stage1"]["s1_circ_legal_issues"]["categoryTitle"],
            "Threshold limb (c): exceptional circumstances or complexity",
        )
        self.assertEqual(
            self.data["stage2"]["s2_care_vulnerable_client"]["categoryTitle"],
            "Care",
        )

    def test_the_fixture_round_trip_is_exact(self):
        expected = {
            "caseDetails": {
                "feeEarnerName": "Jane Doe",
                "matterType": "Public Law Children",
                "caseMatterName": "Re X (Local Authority care proceedings)",
                "courtLevel": "County Court",
            },
            "panelMembership": {}, "stage1": {}, "stage2": {},
            "finalUpliftPercent": "75",
            "evidenceOnFileConfirmed": True,
            "thresholdDeemed": False,
        }
        for key in ("panel_membership_resolution", "panel_membership_children"):
            _block, checkbox = _current_checkbox(key)
            expected["panelMembership"][key] = {
                "checked": True, "label": checkbox["label"]
            }
        for section, keys, explanations in (
            ("stage1", (
                "s1_cse_detailed_knowledge", "s1_cse_marshalling_evidence",
                "s1_circ_legal_issues",
            ), {}),
            ("stage2", (
                "s2_care_vulnerable_client", "s2_resp_no_counsel_drafting",
            ), {
                "s2_care_vulnerable_client": (
                    "Client has diagnosed PTSD and a learning disability assessed "
                    "at borderline range. Conducted instructions across six shorter "
                    "attendances, used plain-language summaries, and brought an "
                    "intermediary to the FHDRA."
                ),
                "s2_resp_no_counsel_drafting": (
                    "Drafted the full position statement, the Scott Schedule, and "
                    "the section 7 response without recourse to Counsel, all settled "
                    "in advance of the case management hearing."
                ),
            }),
        ):
            for key in keys:
                block, checkbox = _current_checkbox(key)
                expected[section][key] = {
                    "checked": True,
                    "label": checkbox["label"],
                    "code": checkbox["code"],
                    "explanation": explanations.get(key, ""),
                    "categoryTitle": block["title"],
                }
        self.assertEqual(self.data, expected)


class TestDeemedRoute(unittest.TestCase):
    """Stage 1 empty + panel + the deemed composite heading, in .docx."""

    @classmethod
    def setUpClass(cls):
        cls.data = extract.extract_formdata(DEEMED)

    def test_deemed_is_affirmatively_extracted(self):
        self.assertEqual(self.data["stage1"], {})
        self.assertTrue(self.data["thresholdDeemed"])

    def test_the_cross_check_passes(self):
        self.assertIsNone(extract.deemed_threshold_support(self.data))

    def test_the_narrative_is_coherent(self):
        self.assertIsNone(extract.threshold_coherence_error(self.data))

    def test_without_the_panel_the_claim_fails(self):
        stripped = dict(self.data, panelMembership={})
        self.assertIsNotNone(extract.deemed_threshold_support(stripped))

    def test_the_fixture_round_trip_is_exact(self):
        panel_block, panel = _current_checkbox("panel_membership_children")
        self.assertEqual(panel_block["page"], 1)
        block, factor = _current_checkbox("s2_resp_no_counsel_drafting")
        self.assertEqual(self.data, {
            "caseDetails": {
                "feeEarnerName": "A. Panel-Member",
                "matterType": "Private Law Children",
                "caseMatterName": "Synthetic Deemed 0001",
                "courtLevel": "County Court",
            },
            "panelMembership": {
                panel["key"]: {"checked": True, "label": panel["label"]},
            },
            "stage1": {},
            "stage2": {
                factor["key"]: {
                    "checked": True,
                    "label": factor["label"],
                    "code": factor["code"],
                    "explanation": (
                        "Drafted every position statement and the final threshold "
                        "document without recourse to counsel throughout the proceedings."
                    ),
                    "categoryTitle": block["title"],
                },
            },
            "finalUpliftPercent": "20",
            "evidenceOnFileConfirmed": False,
            "thresholdDeemed": True,
        })


class TestPasteImmunity(unittest.TestCase):
    """nasty.docx: pasted structure-lookalikes must all stay inert."""

    @classmethod
    def setUpClass(cls):
        cls.data = extract.extract_formdata(NASTY)

    def test_nothing_is_unrecognised(self):
        self.assertNotIn("unrecognised", self.data)

    def test_the_pasted_block_stays_inside_its_explanation(self):
        exp = self.data["stage2"]["s2_complexity_legal_issues"]["explanation"]
        for line in (
            "MATTER DETAIL", "Matter\tFake pasted matter",
            "STAGE 1 : Threshold route", _coded_row("s1_cse_detailed_knowledge"),
            "STAGE 2 : Level of enhancement", _coded_row("s2_care_vulnerable_client"),
            "PROPOSED UPLIFT", "Solicitor’s proposed uplift\t95%",
            "EVIDENCE ON FILE : Confirmed",
        ):
            with self.subTest(line=line):
                self.assertIn(line, exp)

    def test_a_pasted_real_label_creates_no_tick(self):
        """Pasted coded rows are sub-lines, not paragraphs, so they add no ticks."""
        self.assertEqual(list(self.data["stage2"]), ["s2_complexity_legal_issues"])
        self.assertEqual(list(self.data["stage1"]), ["s1_cse_detailed_knowledge"])

    def test_a_pasted_confirmation_cannot_confirm_evidence(self):
        self.assertFalse(self.data["evidenceOnFileConfirmed"])

    def test_a_pasted_percentage_cannot_replace_the_real_one(self):
        self.assertEqual(self.data["finalUpliftPercent"], "40")

    def test_a_matter_name_imitating_the_court_field_stays_whole(self):
        """The exact trap that forced the PDF reader's longest-run logic."""
        self.assertEqual(
            self.data["caseDetails"]["caseMatterName"],
            "In the High Court: Re X and Y (a very long synthetic matter name)",
        )
        self.assertEqual(self.data["caseDetails"]["courtLevel"], "High Court")

    def test_the_control_character_arrived_as_a_space(self):
        exp = self.data["stage2"]["s2_complexity_legal_issues"]["explanation"]
        self.assertIn("The vertical tab above", exp)
        self.assertNotIn("\x0b", exp)

    def test_the_soft_break_heading_prefix_inside_prose_is_not_a_boundary(self):
        """The composite heading is one exact paragraph, prefix plus ``\t\n``
        route. A pasted first line inside an explanation cannot match it."""
        paragraphs = extract_docx.read_docx_paragraphs(NASTY)
        explanation = next(p for p in paragraphs if p.startswith("From my working note"))
        self.assertIn("\nSTAGE 1 : Threshold route\n", explanation)
        self.assertNotEqual(explanation, extract_docx.STAGE1_ESTABLISHED)
        self.assertEqual(extract_docx.structural_damage(paragraphs), [])


class TestFormatDispatch(unittest.TestCase):
    """Content decides the reader; the extension is only a fallback."""

    def test_magic_bytes(self):
        self.assertEqual(extract.detect_format(SAMPLE), "docx")
        self.assertEqual(extract.detect_format(FIXTURES / "sample.pdf"), "pdf")

    def test_a_renamed_docx_still_reads_as_docx(self):
        tmp = Path(tempfile.mkdtemp()) / "misnamed.pdf"
        shutil.copy(SAMPLE, tmp)
        self.assertEqual(extract.detect_format(tmp), "docx")
        data = extract.extract_formdata(tmp)
        self.assertEqual(data["finalUpliftPercent"], "75")

    def test_garbage_falls_back_to_the_extension(self):
        tmp = Path(tempfile.mkdtemp())
        garbage_docx = tmp / "x.docx"
        garbage_docx.write_bytes(b"\x00\x01not a real file")
        self.assertEqual(extract.detect_format(garbage_docx), "docx")
        garbage_pdf = tmp / "x.pdf"
        garbage_pdf.write_bytes(b"\x00\x01not a real file")
        self.assertEqual(extract.detect_format(garbage_pdf), "pdf")


class TestEditedDocuments(unittest.TestCase):
    """A Collator docx only changes by being edited in Word. Every damage
    reading must fail closed — the same discipline as the PDF path."""

    def test_a_cross_section_label_is_refused(self):
        """A Stage 1 label edited into Stage 2 resolves to a key from the
        wrong stage, which must read as damage — never be refiled."""
        s1_label = _current_label("s1_cse_marshalling_evidence")
        edited = _rewrite_paragraph(
            SAMPLE,
            _coded_row("s2_care_vulnerable_client"),
            "CARE 05\t" + s1_label,
        )
        data = extract.extract_formdata(edited)
        self.assertNotIn("s1_cse_marshalling_evidence", data["stage2"])
        unrecognised = data.get("unrecognised") or []
        self.assertTrue(
            any(u["section"] == "stage2" and u["label"] == s1_label
                for u in unrecognised),
            unrecognised,
        )

    def test_a_criterion_label_under_panel_membership_is_refused(self):
        """The costliest misread the PDF path ever had: a criterion accepted
        as a panel membership asserts a guaranteed 15% (CAG 12.20)."""
        s1_label = _current_label("s1_cse_detailed_knowledge")
        edited = _rewrite_paragraph(
            SAMPLE,
            "Memberships\t- Resolution Accredited Specialist Panel",
            "Memberships\t- " + s1_label,
        )
        data = extract.extract_formdata(edited)
        self.assertNotIn("s1_cse_detailed_knowledge", data["panelMembership"])
        self.assertTrue(any(
            u["section"] == "panelMembership" for u in data.get("unrecognised", [])
        ))

    def test_a_damaged_evidence_sentence_reads_as_not_confirmed(self):
        edited = _rewrite_paragraph(
            SAMPLE,
            extract_docx.EVIDENCE_CONFIRMED_SENTENCE,
            "✓   The fee earner says something else entirely.",
        )
        self.assertFalse(extract.extract_formdata(edited)["evidenceOnFileConfirmed"])

    def test_a_negated_evidence_sentence_reads_as_not_confirmed(self):
        """The sentence is matched WHOLE. A prefix match accepted anything
        appended after "…supporting" — including a negation, which made the
        supposedly agreeing pair affirm opposite propositions. Found by
        cross-model review, 7 August 2026."""
        canonical = extract_docx.EVIDENCE_CONFIRMED_SENTENCE
        for damaged in (
            "✓   The fee earner confirms that evidence supporting the matters "
            "set out above is not held on the case file.",
            canonical + " Or so it is claimed.",
        ):
            with self.subTest(damaged=damaged[:60]):
                edited = _rewrite_paragraph(SAMPLE, canonical, damaged)
                self.assertFalse(
                    extract.extract_formdata(edited)["evidenceOnFileConfirmed"]
                )

    def test_a_damaged_status_line_reads_as_not_confirmed(self):
        edited = _rewrite_paragraph(
            SAMPLE,
            extract_docx.EVIDENCE_CONFIRMED_HEADING,
            "EVIDENCE ON FILE : Confirmd",
        )
        self.assertFalse(extract.extract_formdata(edited)["evidenceOnFileConfirmed"])

    def test_the_deemed_line_is_refused_outside_stage_1(self):
        """The composite heading used as Stage 2 prose must not assert deeming."""
        # Replace a Stage 2 explanation with the *entire* valid composite
        # heading. The state machine has already assigned this paragraph to
        # the preceding item, so even exact equality remains inert there.
        explanation = self._sample_explanation()
        edited = _rewrite_paragraph(SAMPLE, explanation, extract_docx.STAGE1_DEEMED)
        self.assertFalse(extract.extract_formdata(edited)["thresholdDeemed"])

    @staticmethod
    def _sample_explanation():
        return (
            "Client has diagnosed PTSD and a learning disability assessed at "
            "borderline range. Conducted instructions across six shorter "
            "attendances, used plain-language summaries, and brought an "
            "intermediary to the FHDRA."
        )

    def test_a_duplicated_heading_stops_extraction_entirely(self):
        """A fake PROPOSED UPLIFT section pasted as real paragraphs above the
        genuine one must not supply the percentage — or anything else. The
        generator writes each heading once, so a duplicate can only mean an
        edited document, and an edited structure is refused whole rather than
        read around. Found by cross-model review, 7 August 2026, which
        reproduced 95 being read instead of the genuine 75."""
        tmp = Path(tempfile.mkdtemp()) / SAMPLE.name
        shutil.copy(SAMPLE, tmp)
        doc = Document(str(tmp))
        # Turn the title into a fake section above the genuine one.
        for index, new_text in ((0, "PROPOSED UPLIFT"),):
            para = doc.paragraphs[index]
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text
        doc.save(str(tmp))
        data = extract.extract_formdata(tmp)
        self.assertEqual(data["finalUpliftPercent"], "")
        self.assertEqual(data["stage1"], {})
        self.assertEqual(data["stage2"], {})
        self.assertFalse(data["evidenceOnFileConfirmed"])
        text = extract.explain_empty_extraction(tmp)
        self.assertIn("appears 2 times", text)
        self.assertIn("re-download", text.lower())

    def test_reordered_headings_stop_extraction_entirely(self):
        """Headings each present once but out of the generator's order can
        also only come from an edit; refused the same way."""
        tmp = Path(tempfile.mkdtemp()) / SAMPLE.name
        shutil.copy(SAMPLE, tmp)
        doc = Document(str(tmp))
        first = next(p for p in doc.paragraphs if p.text == "MATTER DETAIL")
        last = next(
            p for p in doc.paragraphs
            if p.text == extract_docx.EVIDENCE_CONFIRMED_HEADING
        )
        for para, new_text in (
            (first, extract_docx.EVIDENCE_CONFIRMED_HEADING),
            (last, "MATTER DETAIL"),
        ):
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new_text
        doc.save(str(tmp))
        data = extract.extract_formdata(tmp)
        self.assertEqual(data["stage1"], {})
        self.assertEqual(data["stage2"], {})
        self.assertIn("not in the order", extract.explain_empty_extraction(tmp))

    def test_an_edited_label_stops_the_run_with_a_named_nearest(self):
        real = _current_label("s1_cse_marshalling_evidence")
        edited = _rewrite_paragraph(
            SAMPLE, _coded_row("s1_cse_marshalling_evidence"),
            "A03\t" + real.replace("unusual", "unusal"),
        )
        data = extract.extract_formdata(edited)
        self.assertNotIn("s1_cse_marshalling_evidence", data["stage1"])
        match = [u for u in data["unrecognised"] if u["section"] == "stage1"]
        self.assertEqual(len(match), 1)
        self.assertEqual(match[0]["nearest"], real)

    def test_a_code_mismatch_is_structural_damage(self):
        edited = _rewrite_paragraph(
            SAMPLE, _coded_row("s1_cse_marshalling_evidence"),
            "A99\t" + _current_label("s1_cse_marshalling_evidence"),
        )
        data = extract.extract_formdata(edited)
        self.assertEqual(data["stage1"], {})
        diagnostic = extract.diagnose(edited)
        self.assertTrue(any("A99" in problem and "A03" in problem
                            for problem in diagnostic["structural_damage"]))


# Every key diagnose_docx may return. Its output is printed and pasted while
# triaging real, privileged case documents, so it must be structural or
# software names only — never a field value, an explanation, or a client name.
# No filename (it carries the matter name) and no lastModifiedBy (Word writes
# a person's name there) — see the redaction note above diagnose_docx.
ALLOWED_DIAGNOSE_KEYS = {
    "readable", "failure", "paragraphs", "tables", "raw_chars", "creator",
    "resaved_by_another", "created", "modified", "rewritten_after_seconds",
    "made_by_the_app", "structural_damage", "sections",
}
ALLOWED_SECTION_KEYS = {
    "matched", "paragraph_index", "block_paragraphs", "bullet_lines",
    "explanation_markers",
}


class TestDiagnostics(unittest.TestCase):
    def test_diagnose_reports_structure_only(self):
        d = extract.diagnose(SAMPLE)
        self.assertLessEqual(set(d), ALLOWED_DIAGNOSE_KEYS, set(d))
        for name, sec in d["sections"].items():
            with self.subTest(section=name):
                self.assertLessEqual(set(sec), ALLOWED_SECTION_KEYS)

    def test_the_fixture_is_recognised_as_the_apps_own(self):
        d = extract.diagnose(SAMPLE)
        self.assertTrue(d["made_by_the_app"])
        self.assertTrue(d["creator"].startswith("Uplift Collator"))
        self.assertTrue(d["created"])

    def test_every_section_is_found(self):
        d = extract.diagnose(SAMPLE)
        self.assertTrue(all(s.get("matched") for s in d["sections"].values()),
                        d["sections"])

    def test_no_person_name_ever_appears_in_the_diagnostics(self):
        """Word writes the Office account name into lastModifiedBy on save.
        A person's name in "safe to paste anywhere" output is a GDPR leak, so
        only the boolean fact of a re-save survives — and the filename, which
        carries the matter, is not echoed at all. Found by cross-model
        review, 7 August 2026."""
        tmp = Path(tempfile.mkdtemp()) / "Uplift_Justification-Smith 12345.docx"
        shutil.copy(SAMPLE, tmp)
        doc = Document(str(tmp))
        doc.core_properties.last_modified_by = "Priya Solicitor"
        doc.save(str(tmp))
        rendered = json.dumps(extract.diagnose(tmp))
        self.assertNotIn("Priya", rendered)
        self.assertNotIn("Smith 12345", rendered)
        self.assertTrue(extract.diagnose(tmp)["resaved_by_another"])

    def test_a_foreign_creator_name_is_not_echoed(self):
        tmp = Path(tempfile.mkdtemp()) / "foreign-creator.docx"
        doc = Document()
        doc.core_properties.author = "Alex Paralegal"
        doc.add_paragraph("A letter about something else.")
        doc.save(str(tmp))
        d = extract.diagnose(tmp)
        self.assertFalse(d["made_by_the_app"])
        self.assertNotIn("Alex", json.dumps(d))
        self.assertNotIn("Alex", extract.explain_empty_extraction(tmp))

    def test_a_creator_containing_the_app_stamp_plus_a_name_is_not_trusted(self):
        """The redaction must full-match the app's stamp, not find it as a
        substring: "Uplift Collator v1.13 — Priya Solicitor" is not the app's
        pristine stamp, and a substring test would have echoed Priya's name
        through the trusted branch. Round-2 review finding, 7 August 2026."""
        tmp = Path(tempfile.mkdtemp()) / "stamp-plus-name.docx"
        shutil.copy(SAMPLE, tmp)
        doc = Document(str(tmp))
        doc.core_properties.author = "Uplift Collator v1.13 — Priya Solicitor"
        doc.save(str(tmp))
        d = extract.diagnose(tmp)
        self.assertFalse(d["made_by_the_app"])
        self.assertNotIn("Priya", json.dumps(d))

    def test_case_detail_internal_whitespace_survives(self):
        """The single machine tab is removed; user spaces stay case-exact."""
        edited = _rewrite_paragraph(
            SAMPLE, "Fee earner\tJane Doe", "Fee earner\tJane  Q.  Doe"
        )
        details = extract.extract_formdata(edited)["caseDetails"]
        self.assertEqual(details["feeEarnerName"], "Jane  Q.  Doe")

    def test_garbage_is_reported_as_unreadable(self):
        tmp = Path(tempfile.mkdtemp()) / "junk.docx"
        tmp.write_bytes(b"\x00\x01 nothing like a zip")
        d = extract.diagnose(tmp)
        self.assertFalse(d["readable"])
        text = extract.explain_empty_extraction(tmp)
        self.assertIn("could not be read as a Word document", text)
        self.assertIn("Fix:", text)

    def test_a_foreign_docx_is_reported_as_not_the_collators(self):
        tmp = Path(tempfile.mkdtemp()) / "foreign.docx"
        doc = Document()
        doc.add_paragraph("An ordinary letter about something else.")
        doc.save(str(tmp))
        text = extract.explain_empty_extraction(tmp)
        self.assertIn("does not look like an Uplift Collator document", text)
        self.assertIn("Fix:", text)

    def test_no_ticks_is_told_apart_from_no_sections(self):
        edited = _rewrite_paragraph(
            DEEMED,
            _coded_row("s2_resp_no_counsel_drafting"),
            extract_docx.STAGE2_EMPTY,
        )
        # Strip the explanation paragraph too, leaving valid empty sentinels in
        # both Stage sections rather than a damaged/missing section.
        edited2 = _rewrite_paragraph(
            edited,
            "Drafted every position statement and the final threshold document "
            "without recourse to counsel throughout the proceedings.",
            "",
        )
        text = extract.explain_empty_extraction(edited2)
        self.assertIn("no criteria were ticked", text)


class TestEmptyExplanationSentinel(unittest.TestCase):
    """The explanation paragraph is MANDATORY in the grammar: the generator
    prints a fixed sentinel when an entry carries no explanation. The app's
    wizard gate blocks empty explanations, but the generator is also driven
    directly (fixtures, tests, future callers) and cannot assume the gate ran.
    Without the sentinel, the parser's consume-one-paragraph-per-item rule
    would swallow the next item row — found by review before it shipped, so
    pin it forever."""

    def test_the_sentinel_reads_back_as_an_empty_explanation(self):
        formdata = json.loads(
            (FIXTURES / "sample_formdata.json").read_text(encoding="utf-8")
        )
        typed = formdata["stage2"]["s2_care_vulnerable_client"]["explanation"]
        edited = _rewrite_paragraph(
            SAMPLE, typed, extract_docx.EMPTY_EXPLANATION_SENTINEL
        )
        paragraphs = extract_docx.read_docx_paragraphs(edited)
        self.assertEqual(extract_docx.structural_damage(paragraphs), [])
        data = extract.extract_formdata(edited)
        entry = data["stage2"]["s2_care_vulnerable_client"]
        self.assertTrue(entry["checked"])
        self.assertEqual(entry["explanation"], "")

    def test_a_typed_sentinel_lookalike_also_reads_as_empty(self):
        """A solicitor who literally types the sentinel reads back as "" —
        an adjudicated trade-off, not an accident: the typed sentence asserts
        exactly the absence the empty string records, and the generator
        produces byte-identical documents for both (proved in
        tests/test_empty_explanation.js). The first test in this class IS the
        lookalike round trip — the parser cannot and should not distinguish
        the two origins. This one pins the spelling both sides agreed."""
        self.assertEqual(
            extract_docx.EMPTY_EXPLANATION_SENTINEL,
            "No explanation was provided.",
        )


class TestReviewRoundHardening(unittest.TestCase):
    """Adversarial-review regressions, 7 August 2026: paragraph shapes the
    generator cannot produce must read as structural damage, never extract
    quietly. Each of these extracted cleanly (and wrongly) before the fix."""

    def setUp(self):
        self.paragraphs = extract_docx.read_docx_paragraphs(SAMPLE)

    def _row(self, prefix: str) -> int:
        for index, text in enumerate(self.paragraphs):
            if text.startswith(prefix):
                return index
        raise AssertionError(f"no paragraph starts with {prefix!r}")

    def test_a_membership_dash_row_after_proceedings_is_damage(self):
        paras = self.paragraphs[:]
        stray = paras.pop(self._row("- "))
        paras.insert(paras.index(next(
            t for t in paras if t.startswith("Proceedings\t"))) + 1, stray)
        self.assertTrue(extract_docx.structural_damage(paras))

    def test_none_recorded_with_a_continuation_is_damage(self):
        paras = self.paragraphs[:]
        paras[self._row("Memberships\t")] = "Memberships\tNone recorded"
        # the "- Law Society Children Panel" continuation row is still present
        self.assertTrue(extract_docx.structural_damage(paras))

    def test_a_duplicated_stage2_row_is_damage(self):
        paras = self.paragraphs[:]
        care = self._row("CARE 05\t")
        explanation = next(
            i for i in range(care + 1, len(paras)) if paras[i] != "")
        paras.insert(explanation + 1, "A different, quietly winning explanation.")
        paras.insert(explanation + 1, paras[care])
        self.assertTrue(extract_docx.structural_damage(paras))

    def test_swapped_stage2_rows_cross_file_explanations_so_are_damage(self):
        paras = self.paragraphs[:]
        care, resp = self._row("CARE 05\t"), self._row("RESP 02\t")
        paras[care], paras[resp] = paras[resp], paras[care]
        self.assertTrue(extract_docx.structural_damage(paras))

    def test_a_double_dot_percentage_is_damage(self):
        paras = self.paragraphs[:]
        row = self._row("Solicitor’s proposed uplift\t")
        paras[row] = "Solicitor’s proposed uplift\t1..2%"
        self.assertTrue(extract_docx.structural_damage(paras))

    def test_an_invented_ceiling_is_damage(self):
        paras = self.paragraphs[:]
        row = self._row("Applicable ceiling for this court (CAG 12.2)\t")
        paras[row] = "Applicable ceiling for this court (CAG 12.2)\t60%"
        self.assertTrue(extract_docx.structural_damage(paras))


class TestFixturesAreCanonical(unittest.TestCase):
    def test_fixtures_exist_and_are_zip_packages(self):
        for fixture in (SAMPLE, DEEMED, NASTY):
            with self.subTest(fixture=fixture.name):
                self.assertTrue(fixture.is_file())
                self.assertEqual(fixture.read_bytes()[:2], b"PK")


if __name__ == "__main__":
    unittest.main()
