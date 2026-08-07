"""Extract formData from the current Uplift Collator Word summary.

``docx-summary.js`` is the contract.  Its machine-readable units are whole
Word paragraphs: fixed headings, tab-delimited rows, and (for Stage 2) one
opaque explanation paragraph after each coded item row.  In particular, an
explanation is never split on ``\n``.  python-docx returns soft line breaks as
newlines *inside* the paragraph, and keeping that paragraph atomic is what
prevents pasted heading-looking prose from becoming document structure.

This is a clean v1.13 contract.  Earlier PDFs remain the responsibility of
``extract.py``; there is deliberately no reader for the retired Word layout.
The returned dictionary has the same shape as ``extract.extract_formdata``.
"""

from __future__ import annotations

import lzma
import re
import zipfile
import zlib
from datetime import datetime
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from lxml.etree import XMLSyntaxError

from templates import load_content_data


TITLE = "Uplift Justification"
MATTER_HEADING = "MATTER DETAIL"
STAGE1_ESTABLISHED = (
    "STAGE 1 : Threshold route\t\n"
    "Established by the Stage 1 selections below (Spec para 6.13)."
)
STAGE1_DEEMED = (
    "STAGE 1 : Threshold route\t\n"
    "Deemed satisfied for the named fee earner’s work by relevant panel "
    "membership (Spec para 7.23(a))."
)
STAGE2_HEADING = (
    "STAGE 2 : Level of enhancement\t\n"
    "Established by the Stage 2 selections below (Spec para 6.15)."
)
UPLIFT_HEADING = "PROPOSED UPLIFT"
EVIDENCE_CONFIRMED_HEADING = "EVIDENCE ON FILE : Confirmed"
EVIDENCE_UNCONFIRMED_HEADING = "EVIDENCE ON FILE : Unconfirmed"
ABOUT_HEADING = "About this summary"

STAGE1_EMPTY = "No Stage 1 threshold factors selected."
STAGE1_DEEMED_EMPTY = (
    "No Stage 1 threshold factors were selected; the deemed route stated "
    "above applies."
)
STAGE2_EMPTY = "No Stage 2 factors were selected."
# Printed by the generator when a Stage 2 entry carries an empty explanation.
# The app's own wizard gate blocks empty explanations from ever reaching a
# download, but the generator is also driven directly (fixture builder, tests,
# any future caller) and cannot assume its caller enforced that gate. The
# explanation paragraph is mandatory in the grammar — the parser consumes
# exactly one paragraph per item row — so absence is expressed as this
# sentinel, which maps back to "". A solicitor who literally types these words
# also reads back as "": a deliberate, tested trade-off — the typed sentence
# asserts exactly the absence the empty string records.
EMPTY_EXPLANATION_SENTINEL = "No explanation was provided."

EVIDENCE_CONFIRMED_SENTENCE = (
    "✓   The fee earner confirms that evidence supporting the matters set out "
    "above is held on the case file."
)
EVIDENCE_UNCONFIRMED_SENTENCE = (
    "✗   The fee earner has not confirmed that supporting evidence is held on "
    "the case file. The narrative will not assert that it is."
)

# Everything a damaged, truncated, re-compressed or encrypted package can
# throw between zipfile, zlib and python-docx: BadZipFile (bad CRC),
# zlib.error (mangled deflate stream), NotImplementedError (compression
# methods stdlib zipfile cannot read, e.g. AES from DLP tooling),
# RuntimeError (zip-level encryption), OSError (not a file at all).
UNREADABLE_PACKAGE = (
    PackageNotFoundError, KeyError, ValueError, XMLSyntaxError,
    zipfile.BadZipFile, zlib.error, lzma.LZMAError, NotImplementedError,
    RuntimeError, OSError,
)

_UPLIFT_LABEL = "Solicitor’s proposed uplift"
_CEILING_LABEL = "Applicable ceiling for this court (CAG 12.2)"
# Mirrors the app's own percentage gate: ASCII digits, at most one dot. The
# generator can never print "1..2%" or "1e2%", so anything looser is damage.
_PERCENT_RE = re.compile(r"^(?P<value>[0-9]{1,3}(?:\.[0-9]{1,2})?)%$")
# The only ceilings the app prints (CAG 12.2 / Spec 7.22): 50% everywhere,
# 100% in the higher courts.
_CEILING_VALUES = {"50", "100"}
_CODE_LIKE = {
    "stage1": re.compile(r"^[A-Z]\d{2}$"),
    "stage2": re.compile(r"^[A-Z]+(?: [A-Z]+)* \d{2}$"),
}

# Kept under the old semantic keys because diagnose_docx and downstream debug
# tooling consume this shape.  Panel membership now lives inside MATTER DETAIL,
# and the old disclaimer is now the bottom-pinned About frame.
SECTION_HEADINGS: dict[str, str] = {
    "case_details": MATTER_HEADING,
    "panel_membership": MATTER_HEADING,
    "stage1": "STAGE 1 : Threshold route",
    "stage2": "STAGE 2 : Level of enhancement",
    "proposed_uplift": UPLIFT_HEADING,
    "evidence_on_file": "EVIDENCE ON FILE",
    "disclaimer": ABOUT_HEADING,
}

_DETAIL_FIELDS: list[tuple[str, str]] = [
    ("caseMatterName", "Matter"),
    ("feeEarnerName", "Fee earner"),
    ("matterType", "Proceedings"),
    ("courtLevel", "Court"),
]
_DETAIL_BY_PRINTED = {printed: key for key, printed in _DETAIL_FIELDS}


def read_docx_paragraphs(path: str | Path) -> list[str]:
    """Return body paragraphs in document order; footer parts are excluded."""
    document = Document(str(path))
    try:
        return [paragraph.text for paragraph in document.paragraphs]
    except AttributeError as exc:
        # A well-formed package whose document part is not WordprocessingML —
        # Word's own "Strict Open XML" save format uses a namespace
        # python-docx cannot bind, and redaction tooling can leave a foreign
        # XML stub. Converted (not added to UNREADABLE_PACKAGE wholesale, so
        # AttributeError from our own bugs still surfaces loudly).
        raise ValueError(f"not WordprocessingML: {exc}") from exc


def _panel_display(label: str) -> str:
    """Mirror docx-summary.js's panel-label transform exactly."""
    return re.sub(
        r"^Fee earner is on\s+(?:the\s+)?", "", label, flags=re.IGNORECASE
    )


def _limb_heading(title: str) -> str:
    match = re.search(r"limb \(([abc])\)\s*[:—-]?\s*(.*)$", title, re.I)
    if not match:
        return title
    return f"Limb ({match.group(1).lower()}) — {match.group(2)}"


@lru_cache(maxsize=1)
def _content_contract() -> dict:
    blocks = load_content_data()["question_blocks"]
    items: dict[str, dict[str, dict]] = {"stage1": {}, "stage2": {}}
    labels: dict[str, dict[str, dict]] = {"stage1": {}, "stage2": {}}
    for block in blocks:
        section = "stage1" if block.get("page") == 2 else (
            "stage2" if block.get("page") == 3 else None
        )
        if not section:
            continue
        for checkbox in block.get("checkboxes", []):
            # A retired checkbox has no code (its code lives in
            # RESERVED_ITEM_CODES) and the app hides it from the form, so it
            # prints in no current document. Skip it: an old document naming
            # a retired code then resolves as unknown, which
            # _resolve_item_row already reads as damage — fail closed, not
            # KeyError.
            if checkbox.get("retired") or "code" not in checkbox:
                continue
            info = {
                "key": checkbox["key"],
                "code": checkbox["code"],
                "label": checkbox["label"],
                "categoryTitle": block["title"],
                "limbHeading": _limb_heading(block["title"]),
            }
            items[section][info["code"]] = info
            labels[section][info["label"]] = info

    panel_by_display: dict[str, dict] = {}
    for block in blocks:
        if block.get("page") == 1 and block.get("id") == "panel":
            for checkbox in block.get("checkboxes", []):
                panel_by_display[_panel_display(checkbox["label"])] = checkbox

    return {
        "by_code": items,
        "by_label": labels,
        "panel_by_display": panel_by_display,
        "limb_headings": {
            info["limbHeading"]
            for info in items["stage1"].values()
        },
    }


def _split_machine_row(text: str) -> tuple[str, str] | None:
    """Split a generator row, rejecting user-created extra delimiters."""
    if text.count("\t") != 1 or "\n" in text:
        return None
    return tuple(text.split("\t", 1))  # type: ignore[return-value]


def _resolve_item_row(text: str, section: str) -> dict:
    """Describe a coded row without ever inspecting explanation sub-lines."""
    split = _split_machine_row(text)
    if not split:
        return {"itemlike": False, "info": None, "label": text, "damage": None, "code": None}
    code, label = split
    contract = _content_contract()
    by_code = contract["by_code"][section]
    by_label = contract["by_label"][section]
    code_info = by_code.get(code)
    label_info = by_label.get(label)
    itemlike = bool(code_info or label_info or _CODE_LIKE[section].match(code))

    if code_info and label_info and code_info["key"] == label_info["key"]:
        return {"itemlike": True, "info": code_info, "label": label, "damage": None, "code": code}

    if label_info:
        # Exact-label fallback identifies the key, then performs the required
        # cross-check against the frozen code printed beside it.
        return {
            "itemlike": True,
            "info": None,
            "label": label,
            "damage": (
                f"the {section} item {label_info['key']!r} is printed with code "
                f"{code!r}, but content-data.js requires {label_info['code']!r}"
            ),
            "code": code,
        }

    if code_info:
        # The code identifies the intended current item, but a changed label is
        # still an unmatched solicitor selection and must not be silently fixed.
        return {"itemlike": True, "info": None, "label": label, "damage": None, "code": code}

    damage = None
    if itemlike:
        damage = f"the {section} item carries unknown code {code!r}"
    return {"itemlike": itemlike, "info": None, "label": label, "damage": damage, "code": code}


def _nonempty(paragraphs: list[str]) -> list[tuple[int, str]]:
    # Generator dividers and the evidence gap are runless paragraphs whose
    # python-docx text is exactly "".  User text is never stripped here.
    return [(index, text) for index, text in enumerate(paragraphs) if text != ""]


def _analyse(paragraphs: list[str]) -> dict:
    """Parse the fixed outer grammar and identify opaque Stage 2 paragraphs."""
    records = _nonempty(paragraphs)
    indexes: dict[str, int] = {}
    problems: list[str] = []
    opaque: set[int] = set()
    stage2_rows: list[dict] = []
    stage2_explanations: list[int] = []

    appearance = any(
        text == TITLE
        or text == MATTER_HEADING
        or text.startswith("STAGE 1 : Threshold route")
        or text.startswith("STAGE 2 : Level of enhancement")
        or text == UPLIFT_HEADING
        or text.startswith("EVIDENCE ON FILE : ")
        or text == ABOUT_HEADING
        for _index, text in records
    )
    if not appearance:
        return {
            "indexes": indexes, "problems": [], "opaque": opaque,
            "stage2_rows": stage2_rows,
            "stage2_explanations": stage2_explanations,
        }

    def positions(predicate, *, exclude_opaque: bool = True) -> list[int]:
        return [
            index for index, text in records
            if (not exclude_opaque or index not in opaque) and predicate(text)
        ]

    def first_after(predicate, after: int) -> int | None:
        return next(
            (index for index, text in records if index > after and predicate(text)),
            None,
        )

    title_positions = positions(lambda text: text == TITLE)
    if title_positions:
        indexes["title"] = title_positions[0]
    matter = first_after(lambda text: text == MATTER_HEADING, indexes.get("title", -1))
    if matter is not None:
        indexes["case_details"] = matter
        indexes["panel_membership"] = matter

    stage1 = first_after(
        lambda text: text in (STAGE1_ESTABLISHED, STAGE1_DEEMED),
        indexes.get("case_details", -1),
    )
    if stage1 is not None:
        indexes["stage1"] = stage1
    stage2 = first_after(
        lambda text: text == STAGE2_HEADING,
        indexes.get("stage1", indexes.get("case_details", -1)),
    )
    if stage2 is not None:
        indexes["stage2"] = stage2

    # Stage 2 is the only user-prose-bearing structural region.  A valid item
    # row consumes exactly the next non-empty paragraph as its explanation,
    # wholly and opaquely, even if that paragraph itself equals a heading.
    if stage2 is not None:
        cursor = next((n for n, (i, _t) in enumerate(records) if i == stage2), -1) + 1
        while 0 <= cursor < len(records):
            index, text = records[cursor]
            if text == UPLIFT_HEADING:
                indexes["proposed_uplift"] = index
                break
            if text == STAGE2_EMPTY:
                cursor += 1
                if cursor < len(records) and records[cursor][1] == UPLIFT_HEADING:
                    indexes["proposed_uplift"] = records[cursor][0]
                else:
                    problems.append("the Stage 2 empty sentinel is not followed by PROPOSED UPLIFT")
                break

            resolved = _resolve_item_row(text, "stage2")
            if not resolved["itemlike"]:
                problems.append(
                    f"unexpected paragraph in Stage 2 at body index {index}"
                )
                cursor += 1
                continue
            stage2_rows.append({"index": index, **resolved})
            if resolved["damage"]:
                problems.append(resolved["damage"])
            cursor += 1
            if cursor >= len(records):
                problems.append(
                    f"the Stage 2 item at body index {index} has no explanation paragraph"
                )
                break
            explanation_index, _explanation = records[cursor]
            opaque.add(explanation_index)
            stage2_explanations.append(explanation_index)
            cursor += 1

        # Strictly ascending, no repeats: the generator sorts Stage 2 codes
        # and never prints one twice. A duplicated row would otherwise let its
        # explanation silently overwrite the original's; a reordering would
        # cross-file explanations between factors.
        seen_codes = [row["code"] for row in stage2_rows if row["code"] is not None]
        if seen_codes != sorted(set(seen_codes)):
            problems.append("the Stage 2 item codes are not in generator order")

    proposed = indexes.get("proposed_uplift")
    if proposed is not None:
        evidence = first_after(
            lambda text: text in (
                EVIDENCE_CONFIRMED_HEADING, EVIDENCE_UNCONFIRMED_HEADING
            ),
            proposed,
        )
        if evidence is not None:
            indexes["evidence_on_file"] = evidence
    evidence = indexes.get("evidence_on_file")
    if evidence is not None:
        about = first_after(lambda text: text == ABOUT_HEADING, evidence)
        if about is not None:
            indexes["disclaimer"] = about

    # Once explanations are masked, every outer heading is mandatory, unique,
    # and ordered.  This catches edits while leaving pasted prose inert.
    heading_specs = [
        ("title", TITLE, lambda text: text == TITLE),
        ("case_details", MATTER_HEADING, lambda text: text == MATTER_HEADING),
        ("stage1", "STAGE 1 : Threshold route", lambda text: text in (STAGE1_ESTABLISHED, STAGE1_DEEMED)),
        ("stage2", "STAGE 2 : Level of enhancement", lambda text: text == STAGE2_HEADING),
        ("proposed_uplift", UPLIFT_HEADING, lambda text: text == UPLIFT_HEADING),
        ("evidence_on_file", "EVIDENCE ON FILE", lambda text: text in (EVIDENCE_CONFIRMED_HEADING, EVIDENCE_UNCONFIRMED_HEADING)),
        ("disclaimer", ABOUT_HEADING, lambda text: text == ABOUT_HEADING),
    ]
    for name, display, predicate in heading_specs:
        found = positions(predicate)
        if len(found) != 1:
            problems.append(f"the {display!r} heading appears {len(found)} times")
        elif name not in indexes:
            indexes[name] = found[0]
            if name == "case_details":
                indexes["panel_membership"] = found[0]

    ordered_names = [
        "title", "case_details", "stage1", "stage2", "proposed_uplift",
        "evidence_on_file", "disclaimer",
    ]
    ordered = [indexes[name] for name in ordered_names if name in indexes]
    if len(ordered) == len(ordered_names) and ordered != sorted(ordered):
        problems.append("the section headings are not in the order the app writes them")
    if records and indexes.get("title") != records[0][0]:
        problems.append("Uplift Justification is not the first body paragraph")

    # Validate the machine-authored regions as well as their boundaries.
    if "case_details" in indexes and "stage1" in indexes:
        block = [
            text for index, text in records
            if indexes["case_details"] < index < indexes["stage1"]
        ]
        # Strict state machine: dash continuations are legal ONLY immediately
        # after a "Memberships\t- …" row. A dash row anywhere else, or any
        # continuation after "None recorded", is a shape the generator cannot
        # produce and must read as damage — the alternative silently drops or
        # invents a panel, the costliest misread this project has had.
        labels = []
        in_membership_run = False
        for text in block:
            row = _split_machine_row(text)
            if row:
                label, value = row
                labels.append(label)
                in_membership_run = label == "Memberships" and value.startswith("- ")
                # The Memberships value has exactly two generator-legal
                # shapes: the literal "None recorded", or "- " plus a
                # non-empty name. "-", "", a no-space dash or a trailing
                # space would silently erase a panel membership — and a
                # silently dropped panel erases a guaranteed 15% entitlement.
                if label == "Memberships" and not (
                    value == "None recorded"
                    or (value.startswith("- ")
                        and value[2:]
                        and value[2:] == value[2:].strip())
                ):
                    problems.append("the Memberships row is malformed")
            elif text.startswith("- ") and in_membership_run:
                # A continuation, still inside the run. An unknown but
                # non-empty panel name is NOT damage — extract_panel surfaces
                # it as unrecognised, which tolerates future label rewording.
                # But an empty name, or whitespace beyond the "- " prefix, is
                # generator-impossible: the generator prints trimmed labels.
                if not text[2:] or text[2:] != text[2:].strip():
                    problems.append("MATTER DETAIL contains an empty membership row")
            else:
                problems.append("MATTER DETAIL contains a paragraph that is not a machine row")
                in_membership_run = False
        expected = ["Matter", "Fee earner", "Memberships", "Proceedings", "Court"]
        if labels != expected:
            problems.append(
                "the MATTER DETAIL rows are not exactly Matter, Fee earner, "
                "Memberships, Proceedings, Court in that order"
            )

    if "stage1" in indexes and "stage2" in indexes:
        block = [
            (index, text) for index, text in records
            if indexes["stage1"] < index < indexes["stage2"]
        ]
        route_deemed = paragraphs[indexes["stage1"]] == STAGE1_DEEMED
        if len(block) == 1 and block[0][1] in (STAGE1_EMPTY, STAGE1_DEEMED_EMPTY):
            expected = STAGE1_DEEMED_EMPTY if route_deemed else STAGE1_EMPTY
            if block[0][1] != expected:
                problems.append("the Stage 1 route and empty sentinel disagree")
        else:
            current_limb = None
            last_code = ""
            for index, text in block:
                if text in _content_contract()["limb_headings"]:
                    current_limb = text
                    continue
                resolved = _resolve_item_row(text, "stage1")
                if not resolved["itemlike"]:
                    problems.append(f"unexpected paragraph in Stage 1 at body index {index}")
                    continue
                if resolved["damage"]:
                    problems.append(resolved["damage"])
                info = resolved["info"]
                # Placement is validated from the RAW printed code, so an
                # edited label cannot exempt its row from the limb check —
                # the code alone names the limb the generator prints it under.
                raw_info = info
                if raw_info is None and resolved["code"] is not None:
                    raw_info = _content_contract()["by_code"]["stage1"].get(
                        resolved["code"]
                    )
                if raw_info and current_limb != raw_info["limbHeading"]:
                    problems.append(
                        f"the Stage 1 item {raw_info['code']} is under the wrong limb heading"
                    )
                # Strictly ascending over the RAW printed codes: the
                # generator sorts codes and never repeats one, so equality (a
                # duplicated row) is damage too, and an edited label cannot
                # exempt its row from the order check.
                raw_code = resolved["code"]
                if raw_code is not None and raw_code <= last_code:
                    problems.append("the Stage 1 item codes are not in generator order")
                if raw_code is not None:
                    last_code = raw_code

    if proposed is not None and evidence is not None:
        uplift_rows = [
            text for index, text in records if proposed < index < evidence
        ]
        expected_labels = [_UPLIFT_LABEL]
        if len(uplift_rows) == 2:
            expected_labels.append(_CEILING_LABEL)
        parsed_labels = []
        for text in uplift_rows:
            row = _split_machine_row(text)
            # "Not Set%" is generator-legal: an empty finalUpliftPercent prints
            # it (the app's gate prevents that at download, but the generator
            # cannot assume its caller is the app). It reads back as "".
            valid_value = (
                _PERCENT_RE.match(row[1]) is not None
                or (row[0] == _UPLIFT_LABEL and row[1] == "Not Set%")
            ) if row else False
            if not valid_value:
                problems.append("PROPOSED UPLIFT contains a malformed percentage row")
                continue
            if row[0] == _CEILING_LABEL and row[1][:-1] not in _CEILING_VALUES:
                problems.append(
                    "the ceiling row is not one of the two CAG 12.2 ceilings"
                )
            parsed_labels.append(row[0])
        if parsed_labels != expected_labels or len(uplift_rows) not in (1, 2):
            problems.append("the uplift and optional ceiling rows are not in generator order")

    if evidence is not None and "disclaimer" in indexes:
        tail = [(i, t) for i, t in records if i > indexes["disclaimer"]]
        # One disclaimer paragraph, then one glyph sentence.  The disclaimer is
        # machine prose but intentionally not echoed into diagnostics.
        glyph = tail[1][1] if len(tail) >= 2 else None
        heading = paragraphs[evidence]
        expected_glyph = (
            EVIDENCE_CONFIRMED_SENTENCE
            if heading == EVIDENCE_CONFIRMED_HEADING
            else EVIDENCE_UNCONFIRMED_SENTENCE
        )
        if len(tail) != 2 or glyph != expected_glyph:
            problems.append("the evidence heading and its glyph sentence do not agree")

    return {
        "indexes": indexes,
        "problems": list(dict.fromkeys(problems)),
        "opaque": opaque,
        "stage2_rows": stage2_rows,
        "stage2_explanations": stage2_explanations,
    }


def _section_indexes(paragraphs: list[str]) -> dict[str, int]:
    return _analyse(paragraphs)["indexes"]


def structural_damage(paragraphs: list[str]) -> list[str]:
    """Return structural faults that make extraction unsafe, or ``[]``."""
    return _analyse(paragraphs)["problems"]


def section_paragraphs(paragraphs: list[str], name: str) -> list[str]:
    """Return a semantic section using the new paragraph boundaries."""
    indexes = _section_indexes(paragraphs)
    if name not in indexes:
        return []
    next_name = {
        "case_details": "stage1",
        "panel_membership": "stage1",
        "stage1": "stage2",
        "stage2": "proposed_uplift",
        "proposed_uplift": "evidence_on_file",
        "evidence_on_file": "disclaimer",
        "disclaimer": None,
    }[name]
    start = indexes[name] + 1
    end = indexes.get(next_name, len(paragraphs)) if next_name else len(paragraphs)
    return paragraphs[start:end]


def extract_case_details(paragraphs: list[str]) -> dict:
    fields = {key: "" for key, _printed in _DETAIL_FIELDS}
    for text in section_paragraphs(paragraphs, "case_details"):
        row = _split_machine_row(text)
        if row and row[0] in _DETAIL_BY_PRINTED:
            fields[_DETAIL_BY_PRINTED[row[0]]] = row[1]
    return fields


def _panel_displays(paragraphs: list[str]) -> list[str]:
    block = section_paragraphs(paragraphs, "panel_membership")
    displays: list[str] = []
    in_memberships = False
    for text in block:
        row = _split_machine_row(text)
        if row and row[0] == "Memberships":
            in_memberships = True
            if row[1] == "None recorded":
                return []
            if row[1].startswith("- "):
                displays.append(row[1][2:])
            continue
        if in_memberships and text.startswith("- "):
            displays.append(text[2:])
            continue
        if row and row[0] == "Proceedings":
            break
    return displays


def extract_panel(
    paragraphs: list[str],
    label_keys: dict[str, str] | None = None,
    unmatched: list[dict] | None = None,
) -> dict:
    """Rebuild panel entries from the Memberships dash list.

    ``label_keys`` remains accepted for API compatibility but the Word contract
    intentionally uses current content-data labels only; legacy aliases belong
    exclusively to the live PDF path.
    """
    from extract import _unrecognised

    by_display = _content_contract()["panel_by_display"]
    result = {}
    known = list(by_display)
    for display in _panel_displays(paragraphs):
        checkbox = by_display.get(display)
        if not checkbox:
            if unmatched is not None:
                unmatched.append(_unrecognised("panelMembership", display, known))
            continue
        result[checkbox["key"]] = {
            "checked": True,
            "label": checkbox["label"],
        }
    return result


def extract_criteria(
    paragraphs: list[str],
    section_name: str,
    label_keys: dict[str, str] | None = None,
    unmatched: list[dict] | None = None,
) -> dict:
    """Rebuild Stage 1/2 entries from coded rows and current content data."""
    from extract import _unrecognised

    if section_name not in ("stage1", "stage2"):
        return {}
    block = [text for text in section_paragraphs(paragraphs, section_name) if text != ""]
    if not block or block[0] in (STAGE1_EMPTY, STAGE1_DEEMED_EMPTY, STAGE2_EMPTY):
        return {}

    known = list(_content_contract()["by_label"][section_name])
    result = {}
    cursor = 0
    while cursor < len(block):
        text = block[cursor]
        if section_name == "stage1" and text in _content_contract()["limb_headings"]:
            cursor += 1
            continue
        resolved = _resolve_item_row(text, section_name)
        if not resolved["itemlike"]:
            cursor += 1
            continue

        explanation = ""
        if section_name == "stage2" and cursor + 1 < len(block):
            # Opaque by construction: do not strip, split, or grammar-match it.
            # The one exception is the generator's own empty-explanation
            # sentinel, which round-trips back to the empty string the
            # solicitor actually submitted.
            explanation = block[cursor + 1]
            if explanation == EMPTY_EXPLANATION_SENTINEL:
                explanation = ""
            cursor += 1
        info = resolved["info"]
        if not info:
            if unmatched is not None:
                unmatched.append(
                    _unrecognised(
                        section_name,
                        resolved["label"],
                        known,
                        categoryTitle="",
                        explanation=explanation,
                    )
                )
            cursor += 1
            continue
        result[info["key"]] = {
            "checked": True,
            "label": info["label"],
            "code": info["code"],
            "explanation": explanation,
            "categoryTitle": info["categoryTitle"],
        }
        cursor += 1
    return result


def extract_threshold_deemed(paragraphs: list[str]) -> bool:
    indexes = _section_indexes(paragraphs)
    return (
        "stage1" in indexes
        and paragraphs[indexes["stage1"]] == STAGE1_DEEMED
    )


def extract_uplift_percent(paragraphs: list[str]) -> str:
    for text in section_paragraphs(paragraphs, "proposed_uplift"):
        row = _split_machine_row(text)
        if row and row[0] == _UPLIFT_LABEL:
            match = _PERCENT_RE.match(row[1])
            return match.group("value") if match else ""
    return ""


def _evidence_state(paragraphs: list[str]) -> bool | None:
    indexes = _section_indexes(paragraphs)
    evidence = indexes.get("evidence_on_file")
    about = indexes.get("disclaimer")
    if evidence is None or about is None:
        return None
    after_about = [text for text in paragraphs[about + 1:] if text != ""]
    if len(after_about) < 2:
        return None
    glyph = after_about[1]
    heading = paragraphs[evidence]
    if heading == EVIDENCE_CONFIRMED_HEADING and glyph == EVIDENCE_CONFIRMED_SENTENCE:
        return True
    if heading == EVIDENCE_UNCONFIRMED_HEADING and glyph == EVIDENCE_UNCONFIRMED_SENTENCE:
        return False
    return None


def extract_evidence_confirmation(paragraphs: list[str]) -> bool:
    """Return True only for an agreeing Confirmed heading/glyph pair."""
    return _evidence_state(paragraphs) is True


def _empty_formdata() -> dict:
    return {
        "caseDetails": {key: "" for key, _printed in _DETAIL_FIELDS},
        "panelMembership": {},
        "stage1": {},
        "stage2": {},
        "finalUpliftPercent": "",
        "evidenceOnFileConfirmed": False,
        "thresholdDeemed": False,
    }


def extract_formdata_docx(path: str | Path) -> dict:
    try:
        paragraphs = read_docx_paragraphs(path)
    except UNREADABLE_PACKAGE:
        # A corrupt package (truncated transfer, mangled mail gateway,
        # re-compressed or encrypted by intermediate tooling) fails closed
        # exactly like structural damage; diagnose_docx explains it.
        return _empty_formdata()
    if structural_damage(paragraphs):
        return _empty_formdata()
    unmatched: list[dict] = []
    data = {
        "caseDetails": extract_case_details(paragraphs),
        "panelMembership": extract_panel(paragraphs, unmatched=unmatched),
        "stage1": extract_criteria(paragraphs, "stage1", unmatched=unmatched),
        "stage2": extract_criteria(paragraphs, "stage2", unmatched=unmatched),
        "finalUpliftPercent": extract_uplift_percent(paragraphs),
        "evidenceOnFileConfirmed": extract_evidence_confirmation(paragraphs),
        "thresholdDeemed": extract_threshold_deemed(paragraphs),
    }
    if unmatched:
        data["unrecognised"] = unmatched
    return data


# Structural-only diagnostics: values, explanations, filenames and personal
# metadata are deliberately never returned.
_APP_CREATOR_RE = re.compile(r"^Uplift Collator v[\d.]+$")


def _iso(moment: datetime | None) -> str:
    return moment.isoformat() if moment else ""


def diagnose_docx(path: str | Path) -> dict:
    path = Path(path)
    if not zipfile.is_zipfile(path):
        return {"readable": False, "failure": "not_a_zip"}
    try:
        doc = Document(str(path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    except AttributeError:
        # See read_docx_paragraphs: Strict OOXML or a foreign-XML stub.
        return {
            "readable": False,
            "failure": "not_a_word_package: not_wordprocessingml",
        }
    except UNREADABLE_PACKAGE as exc:
        return {
            "readable": False,
            "failure": f"not_a_word_package: {type(exc).__name__}",
        }

    analysis = _analyse(paragraphs)
    indexes = analysis["indexes"]
    # Metadata reads must not decide readability (a blanked
    # <dcterms:created/> from scrubbing tooling raises TypeError deep in
    # python-docx), and a failed read must not discard the reads that
    # already succeeded — the creator stamp is provenance narrate --debug
    # relies on. Defaults are hoisted so each successful read sticks.
    creator = last_modified_by = ""
    created = modified = None
    try:
        cp = doc.core_properties
        creator = cp.author or ""
        last_modified_by = cp.last_modified_by or ""
        created, modified = cp.created, cp.modified
    except Exception:
        pass
    made_by_the_app = bool(_APP_CREATOR_RE.fullmatch(creator.strip()))

    sections = {}
    for name in SECTION_HEADINGS:
        if name not in indexes:
            sections[name] = {"matched": False}
            continue
        block = section_paragraphs(paragraphs, name)
        item_rows = sum(
            1 for text in block
            if _resolve_item_row(text, name)["itemlike"]
        ) if name in ("stage1", "stage2") else (
            len(_panel_displays(paragraphs)) if name == "panel_membership" else 0
        )
        sections[name] = {
            "matched": True,
            "paragraph_index": indexes[name],
            "block_paragraphs": len(block),
            # Names retained for diagnostic-shape compatibility.  Under the
            # new contract these count coded/dash item rows and opaque Stage 2
            # explanation paragraphs, respectively.
            "bullet_lines": item_rows,
            "explanation_markers": (
                len(analysis["stage2_explanations"]) if name == "stage2" else 0
            ),
        }

    return {
        "readable": True,
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "raw_chars": sum(len(text) for text in paragraphs),
        "creator": creator if made_by_the_app else "",
        "made_by_the_app": made_by_the_app,
        "resaved_by_another": bool(last_modified_by and last_modified_by != creator),
        "created": _iso(created),
        "modified": _iso(modified),
        "rewritten_after_seconds": (
            int((modified - created).total_seconds())
            if created and modified else None
        ),
        "structural_damage": analysis["problems"],
        "sections": sections,
    }


def explain_empty_extraction_docx(path: str | Path) -> str:
    """Explain an empty Word extraction without exposing case content."""
    diagnostic = diagnose_docx(path)
    if not diagnostic.get("readable"):
        return (
            "This file could not be read as a Word document at all.\n"
            f"(Technical reason: {diagnostic.get('failure', 'unknown')}.)\n\n"
            "The Collator's own download is a valid .docx, so this happened to\n"
            "the file after the app saved it — an upload, a mail gateway, or a\n"
            "rename of something that was never a Word file.\n\n"
            "Fix: ask for the file exactly as the app saved it, straight from "
            "the browser's Downloads folder."
        )

    if diagnostic.get("structural_damage"):
        problems = "\n".join(
            f"  - {problem}" for problem in diagnostic["structural_damage"]
        )
        return (
            "This document's paragraph structure cannot be trusted, so nothing "
            "was read from it:\n\n"
            f"{problems}\n\n"
            "The current Collator writes one fixed paragraph grammar. This file "
            "does not have that shape, so reading around the damage could treat "
            "edited or pasted text as a solicitor selection.\n\n"
            "Fix: re-download the summary from the current app and send it on "
            "without editing it."
        )

    matched = {
        name for name, section in diagnostic["sections"].items()
        if section.get("matched")
    }
    ticks = sum(
        diagnostic["sections"][name].get("bullet_lines", 0)
        for name in ("stage1", "stage2") if name in matched
    )
    if not matched:
        note = ""
        if not diagnostic.get("made_by_the_app"):
            note = (
                "\nThe document does not carry the Uplift Collator's creator "
                "stamp, so it\nwas made or rebuilt by something else.\n"
            )
        return (
            "This does not look like an Uplift Collator document.\n"
            f"It has text ({diagnostic['raw_chars']} characters in "
            f"{diagnostic['paragraphs']} paragraphs) but none of the current "
            "contract headings.\n"
            + note
            + "\nFix: re-download it from the app's own \"Download Word Summary\" "
            "button."
        )
    if {"stage1", "stage2"} <= matched and ticks == 0:
        return (
            "This is an Uplift Collator document, but no criteria were ticked "
            "in it.\nStage 1 and Stage 2 are both empty, so there is nothing to "
            "justify.\n\nFix: complete the questionnaire, then download it again."
        )
    if not {"stage1", "stage2"} & matched:
        return (
            "This document has Collator sections but no Stage 1 or Stage 2 "
            "section.\n\nFix: re-download it from the current app."
        )
    return (
        f"The Stage sections contain {ticks} item row(s), but none could be "
        "recovered. Look for the unmatched-label report above.\n\n"
        "Fix: re-download the document from the current app without editing its "
        "codes or labels."
    )
