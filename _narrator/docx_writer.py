"""Render the polished narrative as a Word document.

The .docx is the file that actually gets sent. It is **not** a standalone
document: it is the final section of the firm's CCMS bill narrative, pasted in
beneath everything else, so it carries no title block, no page numbers and no
footer — all of which would be wrong once it is inside a longer document. Its
typography matches ``CCMS_BN v13e.docx`` (Arial Nova Cond Light 12pt, single
spaced, A4, 0.5in margins) so the join is invisible.

The house style authority for the surrounding document is
``~/coding/LLM-benchmarking/docs/real-format-notes.md`` §B7, which places this
artifact as section 6, "Claim For Additional General Enhancement". The layout
here was chosen against that document on 1 August 2026:

* the criteria are **bulleted**, because they are a list of discrete grounds
  and an assessor works down them one at a time, striking any that are
  disallowed — prose would make them hunt for where one ground ends;
* headings are **bold only**. Underline conventionally signals emphasis or a
  citation in legal drafting, so using it for headings competes with the
  citations that are the entire point of this section; a full italic line reads
  as quotation. Bold alone also matches the bill narrative's own idiom.

## The rule that matters

``citation-check.txt`` certifies the *Markdown*. If the .docx that gets sent
were missing a citation the check had just certified, the audit record would
attest to a document nobody sent. So this module **never drops prose**, and
:func:`write_docx` re-runs the citation extractor over the finished Word
document and refuses to return if the citations are not identical, in the same
counts, to the Markdown it came from.

That is not theoretical. The model emits headings that carry citations —
``## Threshold Test (Qualifying for Enhancement - Spec Para 6.13 / CAG Section
12.4)`` — so deleting the headings to make the fragment read well inside the
bill narrative, as was done by hand before this module existed, silently took
two certified citations with it. Headings are therefore *demoted* to bold
sub-headings rather than removed, and only the purely structural ones are
dropped at all.
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, Twips

import checks

# ── Typography, taken from CCMS_BN v13e.docx ────────────────────────────────
# Page size is given in twips, exactly as that file stores it. Word rounds page
# dimensions to whole twips on save, so an inch-converted A4 does not survive
# the round trip unchanged; these are the destination document's own numbers.
BODY_FONT = "Arial Nova Cond Light"
BODY_SIZE = Pt(12)
PAGE_WIDTH = Twips(11906)    # A4
PAGE_HEIGHT = Twips(16838)
MARGIN = Inches(0.5)         # 720 twips — exact, no rounding

# The heading the bill narrative already uses for this section, so the fragment
# drops in underneath an identical one rather than needing a rename by hand.
SECTION_HEADING = "Claim For Additional General Enhancement"

# Headings that are pure scaffolding: they label a part of the narrative rather
# than asserting anything, and inside the bill narrative they read as clutter.
# Dropped only after :func:`_heading_is_droppable` confirms the text carries no
# citation — a heading that says anything substantive is demoted, never cut.
STRUCTURAL_HEADINGS = frozenset({"introduction", "conclusion"})


class NarrativeConversionError(Exception):
    """The Word document does not faithfully carry the Markdown it came from.

    Raised rather than returning a document that has quietly lost a citation.
    The Markdown is always written first, so a failure here costs the .docx,
    never the narrative itself.
    """


# ── Block grammar ───────────────────────────────────────────────────────────
# Deliberately small and closed. Anything unrecognised falls through to a plain
# paragraph, which keeps its text: an unknown construct must degrade to prose,
# never to nothing.
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_NUMBERED = re.compile(r"^(\s*)\d+[.)]\s+(.*)$")
_QUOTE = re.compile(r"^>\s?(.*)$")
_RULE = re.compile(r"^\s*([-*_])\s*(?:\1\s*){2,}$")


@dataclass(frozen=True)
class Block:
    kind: str          # heading | bullet | numbered | quote | para
    text: str
    level: int = 0     # heading depth, or list nesting depth


def parse_blocks(markdown: str) -> list[Block]:
    """Split the polished Markdown into renderable blocks.

    Consecutive non-blank lines join into one paragraph, matching Markdown's
    own soft-wrap rule — the model wraps prose across lines and Word must not
    inherit those breaks as hard ones.
    """
    blocks: list[Block] = []
    pending: list[str] = []

    def flush() -> None:
        if pending:
            blocks.append(Block("para", " ".join(pending).strip()))
            pending.clear()

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        if _RULE.match(line):
            # A horizontal rule carries no text, so dropping it loses nothing.
            flush()
            continue
        m = _HEADING.match(line)
        if m:
            flush()
            blocks.append(Block("heading", m.group(2).strip(), len(m.group(1))))
            continue
        m = _BULLET.match(line)
        if m:
            flush()
            blocks.append(Block("bullet", m.group(2).strip(), len(m.group(1)) // 2))
            continue
        m = _NUMBERED.match(line)
        if m:
            flush()
            blocks.append(Block("numbered", m.group(2).strip(), len(m.group(1)) // 2))
            continue
        m = _QUOTE.match(line)
        if m:
            flush()
            blocks.append(Block("quote", m.group(1).strip()))
            continue
        pending.append(line.strip())

    flush()
    return blocks


# ── Inline grammar ──────────────────────────────────────────────────────────
# Bold-italic before bold before italic, so the longer marker wins. Case names
# arrive italicised (*Re B-S (Children) [2013] EWCA Civ 1146*) and that is a
# legal drafting convention worth carrying into Word, not incidental markup.
#
# Italic content may not begin or end with whitespace — Markdown's own rule,
# and the one that stops a stray asterisk in prose ("reviewed 3 * 4 bundles")
# from pairing with a later one and italicising the sentence between them.
_INLINE = re.compile(
    r"\*\*\*(?P<bi>.+?)\*\*\*"
    r"|\*\*(?P<b>.+?)\*\*"
    r"|(?<!\w)\*(?P<i>[^*\s](?:[^*\n]*[^*\s])?)\*(?!\w)"
    r"|`(?P<code>[^`\n]+?)`",
    re.DOTALL,
)


@dataclass(frozen=True)
class Span:
    text: str
    bold: bool = False
    italic: bool = False


def parse_inline(text: str) -> list[Span]:
    """Split a line into styled spans. Unmatched markers stay as literal text."""
    spans: list[Span] = []
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            spans.append(Span(text[pos:m.start()]))
        if m.group("bi") is not None:
            spans.append(Span(m.group("bi"), bold=True, italic=True))
        elif m.group("b") is not None:
            spans.append(Span(m.group("b"), bold=True))
        elif m.group("i") is not None:
            spans.append(Span(m.group("i"), italic=True))
        else:
            spans.append(Span(m.group("code")))
        pos = m.end()
    if pos < len(text):
        spans.append(Span(text[pos:]))
    return [s for s in spans if s.text]


def _heading_is_droppable(text: str) -> bool:
    """True for a heading that is scaffolding and asserts nothing.

    The citation test is the real gate: ``## Threshold Test (Qualifying for
    Enhancement - Spec Para 6.13 / CAG Section 12.4)`` is scaffolding by name
    but carries two authorities, so it must survive as a sub-heading.
    """
    if checks.extract_citations(text):
        return False
    stripped = re.sub(r"[^a-z ]+", "", text.lower()).strip()
    return stripped in STRUCTURAL_HEADINGS


# ── Rendering ───────────────────────────────────────────────────────────────

def _configure(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = section.bottom_margin = MARGIN
        section.left_margin = section.right_margin = MARGIN


def _add(doc: Document, spans: list[Span], *, style: str | None = None,
         bold_all: bool = False, before: int = 0, after: int = 6,
         indent: float | None = None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.0
    if indent is not None:
        para.paragraph_format.left_indent = Inches(indent)
    for span in spans:
        run = para.add_run(span.text)
        run.bold = bold_all or span.bold
        run.italic = span.italic
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
    return para


@dataclass
class ConversionResult:
    """What the conversion did, so the caller can report it honestly."""

    path: Path
    dropped_headings: list[str] = field(default_factory=list)
    demoted_headings: list[str] = field(default_factory=list)
    citations: list[str] = field(default_factory=list)


def build_document(markdown: str) -> tuple[Document, ConversionResult]:
    """Build the Word document in memory. Split out so tests need no disk."""
    doc = Document()
    _configure(doc)
    result = ConversionResult(path=Path())

    _add(doc, [Span(SECTION_HEADING, bold=True)], after=8)

    for block in parse_blocks(markdown):
        spans = parse_inline(block.text)
        if not spans:
            continue

        if block.kind == "heading":
            if _heading_is_droppable(block.text):
                result.dropped_headings.append(block.text)
                continue
            result.demoted_headings.append(block.text)
            _add(doc, spans, bold_all=True, before=8, after=4)

        elif block.kind == "bullet":
            style = "List Bullet" if block.level == 0 else "List Bullet 2"
            _add(doc, spans, style=style, after=4)

        elif block.kind == "numbered":
            _add(doc, spans, style="List Number", after=4)

        elif block.kind == "quote":
            _add(doc, spans, after=4, indent=0.5)

        else:
            _add(doc, spans, after=8)

    return doc, result


def document_text(doc: Document) -> str:
    """Every run of visible text, in reading order.

    Paragraph-joined with newlines so a citation cannot be fabricated by two
    unrelated paragraphs abutting, and run-joined with nothing so a citation
    split across a bold boundary is still found whole.
    """
    return "\n".join("".join(run.text for run in p.runs) for p in doc.paragraphs)


def write_docx(markdown: str, path: str | Path) -> ConversionResult:
    """Render *markdown* to a Word document at *path*.

    Raises :class:`NarrativeConversionError` if the finished document does not
    carry exactly the citations the Markdown did, in the same counts.
    """
    path = Path(path)
    doc, result = build_document(markdown)

    source = Counter(checks._normalise(c) for c in checks.extract_citations(markdown))
    rendered_citations = checks.extract_citations(document_text(doc))
    rendered = Counter(checks._normalise(c) for c in rendered_citations)

    if source != rendered:
        lost = sorted((source - rendered).elements())
        gained = sorted((rendered - source).elements())
        raise NarrativeConversionError(
            "The Word document does not carry the same citations as the "
            "Markdown it came from, so it must not be sent.\n"
            + (f"  Missing from the .docx: {', '.join(lost)}\n" if lost else "")
            + (f"  Appearing only in the .docx: {', '.join(gained)}\n" if gained else "")
            + "The narrative itself is unaffected — narrative-polished.md is "
            "correct and can be used instead."
        )

    result.path = path
    result.citations = rendered_citations
    doc.save(path)
    return result


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <narrative-polished.md> <out.docx>",
              file=sys.stderr)
        sys.exit(1)
    outcome = write_docx(Path(sys.argv[1]).read_text(encoding="utf-8"), sys.argv[2])
    print(f"wrote {outcome.path} — {len(outcome.citations)} citations carried over")
    if outcome.dropped_headings:
        print(f"dropped structural headings: {outcome.dropped_headings}")
